#!/usr/bin/env python3
"""Fusão Fase 2 + Fase 3 no Manual PAR.
Insere 20 blocos (Parte I), a camada dica/exemplo das 14 etapas (Parte II),
reaplica as 2 conformidades CGU-2026 faltantes e anexa os 4 anexos da Fase 3.
Preserva conteúdo e identidade visual; edita o docx via XML."""
import copy, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

U = "/root/.claude/uploads/3517a069-00ce-5e7e-8ba1-0624996120fa"
BASE = U + "/257f12c5-Manual_Calculadora_PAR_atualizado_1.docx"
OUT = "/home/user/desktop-tutorial/Manual_Calculadora_PAR.docx"
SRC = {
 'L1A': U+'/a2cb6d3d-Fase2_Lote1_Blocos_Controversia.docx',
 'L1B': U+'/12e8fba9-Fase2_Lote1_Complemento_Enriquecimento.docx',
 'L2A': U+'/47fb7893-Fase2_Lote2_RitoProcessual.docx',
 'L2B': U+'/cfd85c9e-Fase2_Lote2_Complemento_Enriquecimento.docx',
 'L3':  U+'/79219a5d-Fase2_Lote3_Prescricao_Execucao_Leniencia_Interfaces.docx',
 'PII': U+'/c7ecc279-ParteII_Dosimetria_Etapa_a_Etapa.docx',
 'F31': U+'/65185d26-Fase3_Anexo1_Glossario_Tecnico.docx',
 'F32': U+'/3169701b-Fase3_Anexo2_Indice_Remissivo.docx',
 'F33': U+'/3ea59bc6-Fase3_Anexo3_Quadro_Vigencia_Normativa.docx',
 'F34': U+'/c2094dc9-Fase3_Anexo4_Enunciados_Entendimentos_CGU.docx',
}

base = Document(BASE)
body = base.element.body
W_P, W_TBL = qn('w:p'), qn('w:tbl')
HID = {2: 'Heading2', 3: 'Heading3', 4: 'Heading4', 1: 'Heading1'}

# ── helpers de leitura ────────────────────────────────────────────────────────
def el_text(el):
    return ''.join(t.text or '' for t in el.iter(qn('w:t')))

def is_heading(el):
    if el.tag != W_P: return False
    ppr = el.find(qn('w:pPr'))
    if ppr is None: return False
    st = ppr.find(qn('w:pStyle'))
    return st is not None and (st.get(qn('w:val')) or '').startswith('Heading')

def heading_style_num(el):
    ppr = el.find(qn('w:pPr')); st = ppr.find(qn('w:pStyle'))
    v = st.get(qn('w:val'))
    m = re.match(r'Heading(\d)', v)
    return int(m.group(1)) if m else None

def leading_number(text):
    m = re.match(r'^\s*(\d+(?:\.\d+)*)', text.strip())
    return tuple(int(x) for x in m.group(1).split('.')) if m else None

def children(el):
    return list(el.iterchildren())

# ── região das Partes ─────────────────────────────────────────────────────────
kids = children(body)
def find_part_el(name):
    for el in kids:
        if el.tag == W_P and el_text(el).strip() == name:
            return el
    raise RuntimeError('parte não encontrada: ' + name)
PI_EL, PII_EL = find_part_el('PARTE I'), find_part_el('PARTE II')
pi_i, pii_i = kids.index(PI_EL), kids.index(PII_EL)
REGION_I = kids[pi_i+1:pii_i]
REGION_II = kids[pii_i+1:]

def find_heading_num(num_tuple, region):
    for el in region:
        if is_heading(el):
            n = leading_number(el_text(el))
            if n == num_tuple:
                return el
    raise RuntimeError('heading não encontrado: ' + '.'.join(map(str, num_tuple)))

def find_heading_text(startswith, region):
    for el in region:
        if is_heading(el) and el_text(el).strip().startswith(startswith):
            return el
    raise RuntimeError('heading (texto) não encontrado: ' + startswith)

def boundary_after(anchor_el, anchor_num, region_end_el):
    """Primeiro heading após a seção ancorada (fim da subseção/capítulo)."""
    ta = anchor_num
    el = anchor_el.getnext()
    while el is not None and el is not region_end_el:
        if is_heading(el):
            tn = leading_number(el_text(el))
            if tn is not None:
                if len(ta) == 1:                      # âncora de capítulo
                    if tn[0] > ta[0]:
                        return el
                else:                                  # âncora de subseção
                    if tn[:len(ta)] != ta and tn[0] >= ta[0]:
                        return el
        el = el.getnext()
    return region_end_el

# ── extração de blocos dos arquivos-fonte ─────────────────────────────────────
def load_blocks(path):
    """Retorna dict {num_tuple: [elementos de conteúdo]} para blocos 'Inserção'."""
    d = Document(path)
    ks = children(d.element.body)
    def is_scaffold_h1(el):
        if el.tag != W_P or not is_heading(el) or heading_style_num(el) != 1:
            return False
        t = el_text(el).strip()
        return t.startswith('Inserção') or t.startswith('Próximos passos') or t.startswith('Situação do Lote')
    ins_idx = [i for i, el in enumerate(ks)
               if el.tag == W_P and is_heading(el) and heading_style_num(el) == 1
               and el_text(el).strip().startswith('Inserção')]
    out = {}
    for a in ins_idx:
        b = a + 1; content = []
        while b < len(ks) and not is_scaffold_h1(ks[b]):
            content.append(ks[b]); b += 1
        # primeiro elemento = título H2 do bloco
        title_el = content[0]
        num = leading_number(el_text(title_el))
        out[num] = content
    return out, d   # devolve d para manter árvore viva

# ── inserção de bloco ─────────────────────────────────────────────────────────
def blank_p():
    p = OxmlElement('w:p'); return p

def set_pstyle(p_el, style_id):
    ppr = p_el.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr'); p_el.insert(0, ppr)
    st = ppr.find(qn('w:pStyle'))
    if st is None:
        st = OxmlElement('w:pStyle'); ppr.insert(0, st)
    st.set(qn('w:val'), style_id)

def set_para_text(p_el, text):
    # remove runs existentes, cria um único run com o texto (herda formato do estilo)
    for r in p_el.findall(qn('w:r')):
        p_el.remove(r)
    r = OxmlElement('w:r'); t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p_el.append(r)

def insert_block(boundary_el, content_els, new_num, rename_from=None):
    depth = len(new_num)
    title_style = HID.get(depth, 'Heading2')
    for i, el in enumerate(content_els):
        nel = copy.deepcopy(el)
        if i == 0:
            set_pstyle(nel, title_style)
            if rename_from:
                txt = el_text(nel)
                new_txt = re.sub(r'^\s*' + re.escape(rename_from),
                                 '.'.join(map(str, new_num)), txt, count=1)
                set_para_text(nel, new_txt)
        boundary_el.addprevious(nel)
    boundary_el.addprevious(blank_p())

# ══════════════════════════════════════════════════════════════════════════════
# FASE A — Parte I: 20 blocos
# ══════════════════════════════════════════════════════════════════════════════
# (new_num, source, anchor_existing_num, rename_from)
PLAN = [
 ((1,2,1),   'L1B', (1,2),   None),
 ((2,3,1,1), 'L1A', (2,3,1), None),
 ((2,4),     'L1B', (2,),    '2.5'),
 ((3,1,1),   'L1B', (3,1),   None),
 ((3,2,4),   'L1A', (3,2,3), None),
 ((3,3,3),   'L1A', (3,3,2), None),
 ((4,2,6),   'L1B', (4,2,5), None),
 ((4,3,1),   'L3',  (4,3),   None),
 ((5,2,1,1), 'L2B', (5,2,1), None),
 ((5,2,2,1), 'L2A', (5,2,2), None),
 ((5,2,2,2), 'L2A', (5,2,2), None),
 ((5,3,2,1), 'L2B', (5,3,2), None),
 ((5,3,6,1), 'L2B', (5,3,6), None),
 ((6,2,1,1), 'L2A', (6,2,1), None),
 ((6,5,2,1), 'L2A', (6,5,2), None),
 ((6,7,4),   'L2A', (6,7,3), None),
 ((6,7,5),   'L2B', (6,7,3), None),
 ((8,9),     'L3',  (8,),    '8.7'),
 ((9,5,3),   'L3',  (9,5),   None),
 ((10,2,1),  'L3',  (10,2),  None),
]
srcblocks = {}
_keepalive = []
for key in set(p[1] for p in PLAN):
    blk, d = load_blocks(SRC[key]); srcblocks[key] = blk; _keepalive.append(d)

# Mapear número-do-bloco no arquivo-fonte (pode diferir do destino em caso de rename)
def source_num_for(new_num, rename_from):
    if rename_from:
        return tuple(int(x) for x in rename_from.split('.'))
    return new_num

inserted = 0
# Processar de trás para frente (regra do manifesto): evita que a inserção de um
# irmão vire fronteira do irmão seguinte quando compartilham a mesma âncora.
for new_num, src, anchor_num, rename in reversed(PLAN):
    src_num = source_num_for(new_num, rename)
    content = srcblocks[src].get(src_num)
    assert content, f'bloco-fonte {src_num} não achado em {src}'
    anchor_el = find_heading_num(anchor_num, REGION_I)
    bnd = boundary_after(anchor_el, anchor_num, PII_EL)
    insert_block(bnd, content, new_num, rename)
    inserted += 1
print(f'FASE A: {inserted} blocos da Parte I inseridos.')

# ══════════════════════════════════════════════════════════════════════════════
# FASE B — Parte II: reaplicar conformidades CGU-2026 faltantes
# ══════════════════════════════════════════════════════════════════════════════
CENTER = WD_ALIGN_PARAGRAPH.CENTER
AGRAV, ATEN, CINZA, BRANCO = 'C0392B', '27AE60', 'F2F2F2', 'FFFFFF'

def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    for sh in tcPr.findall(qn('w:shd')): tcPr.remove(sh)
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexfill); tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, align=None, fill=None, size=10):
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    p = cell.paragraphs[0]
    for r in list(p.runs): r._r.getparent().remove(r._r)
    if align is not None: p.alignment = align
    run = p.add_run(text); run.font.name = 'Arial'; run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if fill is not None: shade(cell, fill)

def green_bottom(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    tb = OxmlElement('w:tcBorders'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6'); bot.set(qn('w:space'), '0'); bot.set(qn('w:color'), '64A70B')
    tb.append(bot); tcPr.append(tb)

def new_nota_tbl(text):
    t = base.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.rows[0].cells[0]; set_cell(c, '⚠  ' + text, color='444400', size=10, fill=CINZA); green_bottom(c)
    return t

def new_styled_tbl(data, header_fill, col_widths=None):
    t = base.add_table(rows=len(data), cols=len(data[0])); t.style = 'Table Grid'
    for ri, row in enumerate(data):
        fill = header_fill if ri == 0 else (CINZA if ri % 2 == 1 else BRANCO)
        for ci, val in enumerate(row):
            set_cell(t.rows[ri].cells[ci], val, bold=(ri == 0),
                     color=('FFFFFF' if ri == 0 else None),
                     align=(CENTER if (ri == 0 or ci > 0) else None), fill=fill)
    if col_widths:
        for ri in range(len(data)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)
    return t

def new_body_p(text, bold=False, size=11):
    p = base.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Arial'; r.font.size = Pt(size); r.bold = bold
    return p

# tabelas da Parte II
def tables_in_region(region_els):
    out = []
    for el in region_els:
        if el.tag == W_TBL:
            out.append(el)
    return out

def find_table_by_firstcell(pred):
    from docx.table import Table
    for t in base.tables:
        if t.rows and pred(t):
            return t
    return None

# Ajuste 1 — interrupção 100/400 mil
serv = find_table_by_firstcell(lambda t: any('vila/povoado' in r.cells[0].text for r in t.rows))
assert serv is not None, 'tabela interrupção não achada'
labels1 = [
    'Ausência de interrupção no fornecimento de serviço público',
    'Até 1 semana, ou impacto em município com até 100 mil habitantes',
    'Até 2 semanas, ou impacto em município com até 400 mil habitantes',
    'Até 3 semanas, ou município com mais de 400 mil hab., ou mais de um município do mesmo Estado',
    'Superior a 4 semanas, ou dois ou mais Estados, ou dois ou mais municípios com +400 mil hab.',
]
for ri, txt in enumerate(labels1, start=1):
    set_cell(serv.rows[ri].cells[0], txt, fill=(CINZA if ri % 2 == 1 else BRANCO))
nota_serv = new_nota_tbl('Atualização (CGU, 3ª ed., jun/2026, Tabela 3.1): os limiares de porte de município '
                         'passam a 100 mil e 400 mil habitantes, substituindo os antigos critérios "vila/povoado" e de 500 mil.')
serv._tbl.addnext(nota_serv._tbl)
serv._tbl.addnext(blank_p())

# Ajuste 3 — colaboração faixa discricionária
from docx.table import Table
colab = find_table_by_firstcell(lambda t: 'Condição' in t.rows[0].cells[0].text and 'cumulável' in t.rows[0].cells[0].text)
assert colab is not None, 'tabela colaboração não achada'
while len(colab.rows) > 4:
    colab.rows[-1]._tr.getparent().remove(colab.rows[-1]._tr)
colab_rows = [
    ('Condições factualmente presentes', 'Percentual'),
    ('Nenhuma condição', '0%'),
    ('Uma ou duas condições', 'Faixa discricionária de 0,5% a 1,0% (fixada pela autoridade conforme a utilidade e a relevância da colaboração)'),
    ('As três condições, simultaneamente', '1,5% fixo (sem discricionariedade)'),
]
for ri, (c0, c1) in enumerate(colab_rows):
    if ri == 0:
        set_cell(colab.rows[0].cells[0], c0, bold=True, color='FFFFFF', align=CENTER, fill=ATEN)
        set_cell(colab.rows[0].cells[1], c1, bold=True, color='FFFFFF', align=CENTER, fill=ATEN)
    else:
        fill = CINZA if ri % 2 == 1 else BRANCO
        set_cell(colab.rows[ri].cells[0], c0, fill=fill)
        set_cell(colab.rows[ri].cells[1], c1, align=CENTER, fill=fill)
for ri in range(len(colab.rows)):
    colab.rows[ri].cells[0].width = Cm(5); colab.rows[ri].cells[1].width = Cm(11)
# parágrafo explicativo + atualizar intro/nota
intro_col = None
for el in REGION_II:
    if el.tag == W_P and 'Cada condição acrescenta 0,5%' in el_text(el):
        intro_col = el; break
if intro_col is not None:
    set_para_text(intro_col, 'Reconhecida ainda que não haja admissão de responsabilidade. Conforme a Tabela 6 do '
                  'Manual de Responsabilização da CGU (3ª ed., jun/2026), a colaboração não é soma fixa de 0,5% por '
                  'condição, mas valoração por faixa:')
cond_p = new_body_p('As três condições que se registram (fato objetivo): (i) admitiu a prática do ato; (ii) forneceu '
                    'elementos para a apuração; (iii) renunciou aos prazos processuais. A calculadora deixou de somar '
                    '0,5% por condição e passou a abrir um campo editável (faixa 0,5%–1,0%) quando há uma ou duas '
                    'condições, travando em 1,5% quando há as três.')
colab._tbl.addnext(cond_p._p); colab._tbl.addnext(blank_p())
# nota da colaboração
for t in base.tables:
    if len(t.rows) == 1 and len(t.columns) == 1 and 'mera entrega de documentos' in t.rows[0].cells[0].text:
        set_cell(t.rows[0].cells[0],
                 '⚠  Mudança de metodologia (CGU, 3ª ed., jun/2026): a mera entrega de documentos exigidos por lei NÃO '
                 'configura colaboração — conta a utilidade de informações e provas adicionais e inéditas. Admitir o ato '
                 'sem assumir a responsabilidade jurídica enquadra-se neste inciso (III); o inciso IV exige '
                 'reconhecimento formal da responsabilidade objetiva.', color='444400', size=10, fill=CINZA)
        green_bottom(t.rows[0].cells[0]); break
print('FASE B: conformidades CGU-2026 (Art. 22 III "a" e Art. 23 III) reaplicadas.')

# ══════════════════════════════════════════════════════════════════════════════
# FASE C — Parte II: camada dica prática + exemplo (14 etapas)
# ══════════════════════════════════════════════════════════════════════════════
pii_doc = Document(SRC['PII']); _keepalive.append(pii_doc)
pii_kids = children(pii_doc.element.body)
# extrair, por etapa, apenas as caixas 1x1 "Dica prática de avaliação" e "Exemplo"
etapa_boxes = {}
cur = None
for el in pii_kids:
    if el.tag == W_P and is_heading(el):
        t = el_text(el).strip()
        m = re.match(r'Etapa (\d+)', t)
        cur = int(m.group(1)) if m else (None if heading_style_num(el) == 1 else cur)
    elif el.tag == W_TBL and cur is not None:
        # 1x1?
        trs = el.findall(qn('w:tr'))
        if len(trs) == 1:
            celltxt = el_text(el)
            if celltxt.startswith('Dica prática') or celltxt.startswith('Exemplo'):
                etapa_boxes.setdefault(cur, []).append(el)

# âncora (heading antes do qual inserir) por etapa
ETAPA_BEFORE = {
 0: ('5.',),  1: ('6. AGRAVANTES',),
 2: ('6.2 Art. 22, II',), 3: ('6.3 Art. 22, III',), 4: ('6.4 Art. 22, IV',),
 5: ('6.5 Art. 22, V',), 6: ('6.6 Art. 22, VI',), 7: ('7. ATENUANTES',),
 8: ('7.2 Art. 23, II',), 9: ('7.3 Art. 23, III',), 10: ('7.4 Art. 23, IV',),
 11: ('7.5 Art. 23, V',), 12: ('8. CÁLCULO FINAL',), 13: ('10. RELATÓRIO FINAL',),
}
c_count = 0
for etapa in range(14):
    boxes = etapa_boxes.get(etapa, [])
    if not boxes:
        print('  (aviso) sem caixas para Etapa', etapa); continue
    startswith = ETAPA_BEFORE[etapa][0]
    bnd = find_heading_text(startswith, REGION_II)
    for bx in boxes:
        nb = copy.deepcopy(bx)
        bnd.addprevious(blank_p())
        bnd.addprevious(nb)
        c_count += 1
    bnd.addprevious(blank_p())
print(f'FASE C: {c_count} caixas (dica/exemplo) inseridas nas 14 etapas.')

# ══════════════════════════════════════════════════════════════════════════════
# FASE D — Anexos I–IV (Fase 3), após a Bibliografia
# ══════════════════════════════════════════════════════════════════════════════
def heading1_p(text):
    p = base.add_paragraph(); set_pstyle(p._p, 'Heading1')
    r = p.add_run(text); r.font.name = 'Arial'
    return p

ANEXOS = [
 ('ANEXO I — GLOSSÁRIO TÉCNICO', 'F31'),
 ('ANEXO II — ÍNDICE REMISSIVO ANALÍTICO', 'F32'),
 ('ANEXO III — QUADRO DE VIGÊNCIA NORMATIVA', 'F33'),
 ('ANEXO IV — ENUNCIADOS E ENTENDIMENTOS ADMINISTRATIVOS DA CGU', 'F34'),
]
SECTPR = body.find(qn('w:sectPr'))   # inserir SEMPRE antes do sectPr final
def append_source_content(path, skip_leading_meta=True):
    """Anexa o conteúdo do arquivo-fonte antes do sectPr final,
    pulando os 1-2 parágrafos de cabeçalho 'FASE 3 — ANEXO N' e o subtítulo."""
    d = Document(path); _keepalive.append(d)
    ks = children(d.element.body)
    skipped = 0
    for el in ks:
        if el.tag == qn('w:sectPr'):
            continue  # não copiar o sectPr do arquivo-fonte
        if skip_leading_meta and el.tag == W_P:
            t = el_text(el).strip()
            if skipped < 2 and (t.startswith('FASE 3') or t.startswith('Glossário Técnico') or
                                t.startswith('Índice Remissivo') or t.startswith('Quadro de Vigência') or
                                t.startswith('Enunciados e Entendimentos')):
                skipped += 1; continue
        if el.tag == W_P and el_text(el).strip().startswith('Próximo anexo'):
            continue
        SECTPR.addprevious(copy.deepcopy(el))

for titulo, key in ANEXOS:
    # page break + título do anexo
    pb = base.add_paragraph(); run = pb.add_run(); br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); run._r.append(br)
    heading1_p(titulo)
    append_source_content(SRC[key])
print('FASE D: 4 anexos da Fase 3 adicionados após a Bibliografia.')

# ══════════════════════════════════════════════════════════════════════════════
# FASE E — forçar atualização do sumário (campo TOC) ao abrir no Word
# ══════════════════════════════════════════════════════════════════════════════
settings = base.settings.element
if settings.find(qn('w:updateFields')) is None:
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); settings.append(uf)
print('FASE E: updateFields=true (Word recalcula o sumário na abertura).')

base.save(OUT)
print('\nOK — manual mesclado salvo em', OUT)
