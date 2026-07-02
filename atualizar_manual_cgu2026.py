#!/usr/bin/env python3
"""Aplica ao manual (docx) as 3 modificações do Manual CGU (3ª ed., jun/2026)
já implementadas na calculadora, preservando conteúdo e identidade visual.
Edita diretamente o docx enviado pelo usuário."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = '/root/.claude/uploads/3517a069-00ce-5e7e-8ba1-0624996120fa/03e9eae1-Manual_Calculadora_PAR.docx'
OUT = '/home/user/desktop-tutorial/Manual_Calculadora_PAR.docx'

CENTER = WD_ALIGN_PARAGRAPH.CENTER
AGRAV, ATEN, CINZA, BRANCO = 'C0392B', '27AE60', 'F2F2F2', 'FFFFFF'
doc = Document(SRC)

# ── helpers de estilo (replicam a identidade do documento) ────────────────────
def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    for sh in tcPr.findall(qn('w:shd')):
        tcPr.remove(sh)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexfill)
    tcPr.append(shd)

def clear_para(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)

def set_cell(cell, text, bold=False, color=None, align=None, fill=None, size=10):
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    p = cell.paragraphs[0]
    clear_para(p)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Arial'; run.font.size = Pt(size); run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if fill is not None:
        shade(cell, fill)

def new_styled_table(data, header_fill, col_widths=None):
    ncols = len(data[0])
    t = doc.add_table(rows=len(data), cols=ncols)
    t.style = 'Table Grid'
    for ri, row in enumerate(data):
        fill = header_fill if ri == 0 else (CINZA if ri % 2 == 1 else BRANCO)
        for ci, val in enumerate(row):
            set_cell(t.rows[ri].cells[ci], val, bold=(ri == 0),
                     color=('FFFFFF' if ri == 0 else None),
                     align=(CENTER if (ri == 0 or ci > 0) else None), fill=fill)
    if col_widths:
        for ri in range(len(data)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)
    return t

def new_nota(text):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    set_cell(c, '⚠  ' + text, color='444400', size=10, fill=CINZA)
    tcPr = c._tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6'); bot.set(qn('w:space'), '0'); bot.set(qn('w:color'), '64A70B')
    tb.append(bot); tcPr.append(tb)
    return t

def new_body(runs, style='Normal', size=11):
    p = doc.add_paragraph(); p.style = style
    for txt, bold in runs:
        r = p.add_run(txt); r.font.name = 'Arial'; r.font.size = Pt(size); r.bold = bold
    return p

def new_blank():
    return doc.add_paragraph()

def replace_text_keep_fmt(p, new_text):
    """Substitui o texto do parágrafo preservando o formato do 1º run."""
    runs = p.runs
    if not runs:
        r = p.add_run(new_text); r.font.name = 'Arial'; r.font.size = Pt(11); return
    runs[0].text = new_text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)

def insert_after(ref_el, elements):
    cur = ref_el
    for el in elements:
        cur.addnext(el); cur = el

# ── localizadores ─────────────────────────────────────────────────────────────
def find_para(pred):
    for p in doc.paragraphs:
        if pred(p):
            return p
    raise RuntimeError('parágrafo não encontrado')

def iter_tables():
    return doc.tables

# ═══════════════════════════════════════════════════════════════════════════
# AJUSTE 1 — Art. 22, III, "a": faixas populacionais (Tabela 3.1)
# ═══════════════════════════════════════════════════════════════════════════
serv_tbl = None
for t in iter_tables():
    cells0 = [r.cells[0].text for r in t.rows]
    if any('vila/povoado' in c for c in cells0):
        serv_tbl = t; break
assert serv_tbl is not None, 'tabela de serviço público não encontrada'
labels1 = [
    'Ausência de interrupção no fornecimento de serviço público',
    'Até 1 semana, ou impacto em município com até 100 mil habitantes',
    'Até 2 semanas, ou impacto em município com até 400 mil habitantes',
    'Até 3 semanas, ou município com mais de 400 mil hab., ou mais de um município do mesmo Estado',
    'Superior a 4 semanas, ou dois ou mais Estados, ou dois ou mais municípios com +400 mil hab.',
]
for ri, txt in enumerate(labels1, start=1):
    fill = CINZA if ri % 2 == 1 else BRANCO
    set_cell(serv_tbl.rows[ri].cells[0], txt, fill=fill)
# nota após a tabela de serviço público (inserir antes do parágrafo "b) Interrupção na Execução")
anchor_b = find_para(lambda p: p.text.strip().startswith('b) Interrupção na Execução'))
nota1 = new_nota('Atualização do Manual CGU (3ª ed., jun/2026), Tabela 3.1: os limiares de porte de município '
                 'passam a ser 100 mil e 400 mil habitantes, substituindo os antigos critérios "vila/povoado" e de 500 mil habitantes.')
insert_after(anchor_b._p.getprevious(), [nota1._tbl, new_blank()._p])

# ═══════════════════════════════════════════════════════════════════════════
# AJUSTE 3 — Art. 23, III: colaboração como faixa discricionária (Tabela 6)
# ═══════════════════════════════════════════════════════════════════════════
colab_intro = find_para(lambda p: 'Cada condição acrescenta 0,5%, com teto de 1,5%' in p.text)
replace_text_keep_fmt(colab_intro,
    'Reconhecida ainda que não haja admissão de responsabilidade. Conforme a Tabela 6 do Manual CGU '
    '(3ª ed., jun/2026), a colaboração não é soma fixa de 0,5% por condição, mas valoração por faixa:')

# tabela colab: localizar (header "Condição (cumulável)")
colab_tbl = None
for t in iter_tables():
    if t.rows and 'Condição' in t.rows[0].cells[0].text and 'cumulável' in t.rows[0].cells[0].text:
        colab_tbl = t; break
assert colab_tbl is not None, 'tabela de colaboração não encontrada'
# reduzir de 5 para 4 linhas (header + 3)
while len(colab_tbl.rows) > 4:
    last = colab_tbl.rows[-1]
    last._tr.getparent().remove(last._tr)
colab_rows = [
    ('Condições factualmente presentes', 'Percentual'),
    ('Nenhuma condição', '0%'),
    ('Uma ou duas condições', 'Faixa discricionária de 0,5% a 1,0% (fixada pela autoridade conforme a utilidade e a relevância da colaboração)'),
    ('As três condições, simultaneamente', '1,5% fixo (sem discricionariedade)'),
]
for ri, (c0, c1) in enumerate(colab_rows):
    if ri == 0:
        set_cell(colab_tbl.rows[0].cells[0], c0, bold=True, color='FFFFFF', align=CENTER, fill=ATEN)
        set_cell(colab_tbl.rows[0].cells[1], c1, bold=True, color='FFFFFF', align=CENTER, fill=ATEN)
    else:
        fill = CINZA if ri % 2 == 1 else BRANCO
        set_cell(colab_tbl.rows[ri].cells[0], c0, fill=fill)
        set_cell(colab_tbl.rows[ri].cells[1], c1, align=CENTER, fill=fill)
# larguras
for ri in range(len(colab_tbl.rows)):
    colab_tbl.rows[ri].cells[0].width = Cm(5)
    colab_tbl.rows[ri].cells[1].width = Cm(11)

# parágrafo explicativo (condições factuais + comportamento da calculadora), inserir após a tabela colab
cond_p = new_body([
    ('As três condições que se registram (fato objetivo): (i) admitiu a prática do ato; (ii) forneceu elementos '
     'para a apuração; (iii) renunciou aos prazos processuais. A calculadora deixou de somar 0,5% por condição e '
     'passou a abrir um campo editável (faixa 0,5%–1,0%) quando há uma ou duas condições, travando em 1,5% quando '
     'há as três.', False)
])
insert_after(colab_tbl._tbl, [new_blank()._p, cond_p._p])

# atualizar a nota (⚠) da colaboração
colab_nota_cell = None
for t in iter_tables():
    if len(t.rows) == 1 and len(t.columns) == 1 and 'mera entrega de documentos' in t.rows[0].cells[0].text:
        colab_nota_cell = t.rows[0].cells[0]; break
assert colab_nota_cell is not None, 'nota de colaboração não encontrada'
set_cell(colab_nota_cell,
    '⚠  Mudança de metodologia (Manual CGU 3ª ed., jun/2026): a mera entrega de documentos exigidos por lei NÃO '
    'configura colaboração — conta a utilidade de informações e provas adicionais e inéditas. Admitir o ato sem '
    'assumir a responsabilidade jurídica enquadra-se neste inciso (III); o inciso IV exige reconhecimento formal '
    'da responsabilidade objetiva.',
    color='444400', size=10, fill=CINZA)
# restaurar borda verde inferior da nota
tcPr = colab_nota_cell._tc.get_or_add_tcPr()
for old in tcPr.findall(qn('w:tcBorders')):
    tcPr.remove(old)
tb = OxmlElement('w:tcBorders'); bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6'); bot.set(qn('w:space'), '0'); bot.set(qn('w:color'), '64A70B')
tb.append(bot); tcPr.append(tb)

# ═══════════════════════════════════════════════════════════════════════════
# AJUSTE 2 — Art. 23, V: fator multiplicador 1,25 (Portaria Conjunta CGU 6/2022)
# ═══════════════════════════════════════════════════════════════════════════
formula_p = find_para(lambda p: p.text.strip().startswith('Fórmula: [COI × MPI] + APJ'))
replace_text_keep_fmt(formula_p,
    'Resultado da planilha = [COI × MPI] + APJ   →   se ≥ 1%, aplica-se × 1,25   →   Atenuante = mín.(resultado, 5%)')

onde_p = find_para(lambda p: p.text.strip().startswith('Onde: COI ='))
fator_p = new_body([
    ('Fator multiplicador de 1,25 (Portaria Conjunta CGU nº 6/2022 — Adendo nº 1 ao Manual Prático de Avaliação de PI): ', True),
    ('incide sobre o resultado da planilha sempre que este for igual ou superior a 1%. O teto de 5% é aplicado '
     'DEPOIS do multiplicador.', False),
])
fator_tbl = new_styled_table(
    [
        ('Resultado da planilha', 'Fator 1,25', 'Resultado final'),
        ('Inferior a 1% (ex.: 0,8%)', 'Não incide', '0,8%'),
        ('Igual ou superior a 1% (ex.: 1,0%)', 'Incide', '1,25%'),
        ('Próximo ao teto (ex.: 4,2%)', 'Incide, mas o teto prevalece', '5,0% (4,2 × 1,25 = 5,25 → teto)'),
    ],
    header_fill=ATEN, col_widths=[6, 5, 6])
fator_nota = new_nota('A calculadora exibe o valor pré e pós-multiplicador no painel, no relatório e na minuta, '
                      'para transparência do cálculo.')
insert_after(onde_p._p, [new_blank()._p, fator_p._p, new_blank()._p,
                         fator_tbl._tbl, new_blank()._p, fator_nota._tbl, new_blank()._p])

# ═══════════════════════════════════════════════════════════════════════════
# Referências complementares (2.3): incluir Manual CGU 2026 e Portaria 6/2022
# ═══════════════════════════════════════════════════════════════════════════
ref_in = find_para(lambda p: p.style.name == 'List Bullet' and 'IN CGU nº 1/2015' in p.text)
b1 = new_body([('Manual de Responsabilização de Entes Privados da CGU (3ª ed., jun/2026) — metodologia oficial de '
                'dosimetria (Tabela 3.1: interrupção de serviço público; Tabela 6: colaboração);', False)],
               style='List Bullet')
b2 = new_body([('Portaria Conjunta CGU nº 6/2022 — Adendo nº 1 ao Manual Prático de Avaliação de Programa de '
                'Integridade; fator multiplicador de 1,25 (Art. 23, V);', False)],
               style='List Bullet')
insert_after(ref_in._p, [b1._p, b2._p])

# ═══════════════════════════════════════════════════════════════════════════
# FAQ: duas novas perguntas (após "Sim, mas com percentual reduzido...")
# ═══════════════════════════════════════════════════════════════════════════
faq_anchor = find_para(lambda p: p.text.strip().startswith('Sim, mas com percentual reduzido. O bloco APJ-Posterior'))
q1 = new_body([('Como a colaboração é valorada com 1, 2 ou 3 condições presentes? (atualização CGU 2026)', True)], style='Heading 2', size=11.5)
a1 = new_body([('Conforme a Tabela 6 do Manual CGU, uma ou duas condições ensejam percentual discricionário de 0,5% a '
                '1,0% (fixado pela autoridade conforme a utilidade da colaboração); as três condições, simultaneamente, '
                'ensejam 1,5% fixo. A calculadora abre um campo editável (faixa 0,5–1,0%) para 1 ou 2 condições e trava '
                'em 1,5% para 3 — não soma mais 0,5% por condição.', False)])
q2 = new_body([('O que é o fator de 1,25 no Programa de Integridade? (atualização CGU 2026)', True)], style='Heading 2', size=11.5)
a2 = new_body([('Conforme a Portaria Conjunta CGU nº 6/2022, ao resultado da planilha [COI × MPI] + APJ aplica-se um '
                'fator multiplicador de 1,25 sempre que esse resultado for igual ou superior a 1%, respeitado o teto de '
                '5%. O fator não incide sobre resultados abaixo de 1%, e o teto de 5% é aplicado depois do multiplicador '
                '(ex.: 4,2% × 1,25 = 5,25% → limitado a 5,0%).', False)])
insert_after(faq_anchor._p, [q1._p, a1._p, q2._p, a2._p])

# ═══════════════════════════════════════════════════════════════════════════
# Rodapé: atualizar linha de fundamentos
# ═══════════════════════════════════════════════════════════════════════════
footer_p = find_para(lambda p: 'Manual elaborado com base na Calculadora' in p.text)
# o parágrafo tem 1 run com \n; substituir preservando formatação
if footer_p.runs:
    footer_p.runs[0].text = ('Manual elaborado com base na Calculadora de Sanções — PAR — Corregedoria da Receita '
                             'Federal do Brasil\nFundamentos: Lei nº 12.846/2013 | Decreto nº 11.129/2022 | IN CGU nº '
                             '1/2015 | Manual CGU (3ª ed., jun/2026) | Portaria Conjunta CGU nº 6/2022')
    for r in footer_p.runs[1:]:
        r._r.getparent().remove(r._r)

doc.save(OUT)
print('OK — manual atualizado salvo em', OUT)
