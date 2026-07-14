# -*- coding: utf-8 -*-
"""
COGER Design System — helpers de geração de documentos Word via python-docx.
Implementação fiel ao guia "COGER Design System — Guia de Construção de
Documentos Word" (Corregedoria da Receita Federal do Brasil).
"""
import re

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# 1. Paleta de cores
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GOLD = RGBColor(0xBF, 0xA0, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED_C = RGBColor(0x70, 0x70, 0x70)
BLUE_D = RGBColor(0x2E, 0x75, 0xB6)
GREEN = RGBColor(0x1A, 0x6E, 0x1A)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN_B = "E8F5E9"


def rgb_hex(color):
    """RGBColor -> string hex sem '#', maiúscula."""
    return "%02X%02X%02X" % (color[0], color[1], color[2])


# ---------------------------------------------------------------------------
# 8. Helpers XML de baixo nível
# ---------------------------------------------------------------------------
def set_shading_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    pPr.append(shd)


def set_shading_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def set_left_border(para, color, sz=18):
    """sz em oitavos de ponto: 18 = 2.25pt | 24 = 3pt."""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), rgb_hex(color) if isinstance(color, RGBColor) else str(color))
    pBdr.append(left)


def set_bottom_border(para, color, sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), rgb_hex(color) if isinstance(color, RGBColor) else str(color))
    pBdr.append(bottom)


def set_cell_borders(cell, color="CCCCCC"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _remove_numbering(style):
    """Remove w:numPr do pPr do estilo, evitando numeração automática indesejada."""
    pPr = style.element.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


# ---------------------------------------------------------------------------
# 3. Estilos de heading nativos (compatíveis com TOC automático)
# ---------------------------------------------------------------------------
def setup_heading_styles(doc):
    styles = doc.styles

    h1 = styles["Heading 1"]
    h1.font.name = "Arial Narrow"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(4)
    h1.paragraph_format.space_after = Pt(12)
    _remove_numbering(h1)

    h2 = styles["Heading 2"]
    h2.font.name = "Arial Narrow"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(5)
    _remove_numbering(h2)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(3)
    _remove_numbering(h3)


# ---------------------------------------------------------------------------
# 4. Sumário automático (campo Word TOC)
# ---------------------------------------------------------------------------
def insert_toc(doc, levels="1-3"):
    p = doc.add_paragraph()
    run = p.add_run()

    fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin")
    run._r.append(fb)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "%s" \\h \\z \\u ' % levels
    run._r.append(instr)

    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    run._r.append(fs)
    ph = OxmlElement("w:t")
    ph.text = "Clique com o botão direito → Atualizar campo para ver o sumário com páginas."
    run._r.append(ph)

    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    run._r.append(fe)
    return p


# ---------------------------------------------------------------------------
# 7. Helpers de formatação inline
# ---------------------------------------------------------------------------
def add_runs(para, text, size=11, font="Calibri", color=None, base_bold=False, base_italic=False):
    parts = re.split(r"(\*\*|\*)", text)
    bold = base_bold
    italic = base_italic
    for seg in parts:
        if seg == "**":
            bold = not bold
            continue
        if seg == "*":
            italic = not italic
            continue
        if not seg:
            continue
        r = para.add_run(seg)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color


def add_runs_from_source(para, src_para, size=11, font="Calibri", color=None):
    for src_r in src_para.runs:
        if not src_r.text:
            continue
        r = para.add_run(src_r.text)
        r.font.name = font
        r.font.size = Pt(size)
        r.font.bold = src_r.bold
        r.font.italic = src_r.italic
        if color:
            r.font.color.rgb = color


# ---------------------------------------------------------------------------
# 5. Componentes visuais
# ---------------------------------------------------------------------------
def add_chapter(doc, eyebrow_text, title):
    p_eye = doc.add_paragraph()
    p_eye.paragraph_format.space_after = Pt(2)
    r = p_eye.add_run(eyebrow_text.upper())
    r.font.name = "Arial Narrow"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = GOLD

    p_h = doc.add_paragraph(style="Heading 1")
    r2 = p_h.add_run(title)
    r2.font.name = "Arial Narrow"
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = NAVY

    add_separator(doc, color=GOLD, sz=8)


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    set_left_border(p, GOLD, 18)
    r = p.add_run(text)
    r.font.name = "Arial Narrow"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text, size=size)
    return p


def add_bullet(doc, text, indent=0.7):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(13.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("— ")
    add_runs(p, text)
    return p


def add_numbered(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = Pt(13.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p.add_run("%s.  " % num)
    r0.font.bold = True
    r0.font.color.rgb = NAVY
    r0.font.name = "Calibri"
    r0.font.size = Pt(11)
    add_runs(p, text)
    return p


_ALERT_BG = {"info": "EBF3FB", "warn": "FFF8E1", "danger": "FDE8E8", "green": "E8F5E9", "mono": "F4F4F4"}
_ALERT_BORDER = {"info": BLUE_D, "warn": GOLD, "danger": RED, "green": GREEN, "mono": MUTED_C}


def add_alert(doc, lines, kind="info", label=None):
    bg = _ALERT_BG[kind]
    bdc = _ALERT_BORDER[kind]
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(5) if i == 0 else Pt(0)
        p.paragraph_format.space_after = Pt(3) if i == len(lines) - 1 else Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_left_border(p, bdc, sz=24)
        set_shading_para(p, bg)
        if i == 0 and label:
            rl = p.add_run(label.upper() + "  ")
            rl.font.name = "Arial Narrow"
            rl.font.size = Pt(9)
            rl.font.bold = True
            rl.font.color.rgb = bdc
        add_runs(p, line, size=10, font="Calibri")
    return None


def add_separator(doc, color=None, sz=4):
    color = color or GOLD
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    set_bottom_border(p, color, sz=sz)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def add_muted(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = MUTED_C
    return p


# ---------------------------------------------------------------------------
# 6. Tabelas
# ---------------------------------------------------------------------------
def make_table(doc, headers, rows, col_widths, alt_bg="F5F8FF", header_bg=None):
    header_bg = header_bg or rgb_hex(NAVY)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # largura das colunas (header + cada linha, todas as células)
    def _apply_widths(row):
        for cell, w in zip(row.cells, col_widths):
            cell.width = Cm(w)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = WHITE
        set_shading_cell(cell, header_bg)
        set_cell_borders(cell)
    _apply_widths(table.rows[0])

    for ridx, row_vals in enumerate(rows):
        row = table.add_row()
        bg = alt_bg if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row_vals):
            cell = row.cells[cidx]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Calibri"
            r.font.size = Pt(9.5)
            set_shading_cell(cell, bg)
            set_cell_borders(cell)
        _apply_widths(row)

    return table


# ---------------------------------------------------------------------------
# Imagem com legenda numerada
# ---------------------------------------------------------------------------
_FIGURE_COUNTER = {"n": 0}


def add_figure(doc, image_path, caption, width_cm=15.5):
    _FIGURE_COUNTER["n"] += 1
    n = _FIGURE_COUNTER["n"]
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.add_run().add_picture(image_path, width=Cm(width_cm))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(10)
    r = p_cap.add_run("Figura %d — %s" % (n, caption))
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = MUTED_C
    return n


def reset_figure_counter():
    _FIGURE_COUNTER["n"] = 0


# ---------------------------------------------------------------------------
# 2. Configuração de página / documento base
# ---------------------------------------------------------------------------
def new_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].paragraph_format.space_after = Pt(5)

    setup_heading_styles(doc)
    return doc


# ---------------------------------------------------------------------------
# 5.1 Capa
# ---------------------------------------------------------------------------
def add_cover(doc, logo_path, titulo, subtitulo, metadados_lines, alerta_lines=None):
    p_logo = doc.add_paragraph()
    p_logo.paragraph_format.space_before = Pt(20)
    p_logo.paragraph_format.space_after = Pt(14)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path:
        p_logo.add_run().add_picture(logo_path, width=Cm(7))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run(titulo)
    r_title.font.name = "Arial Narrow"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(10)
    r_sub = p_sub.add_run(subtitulo)
    r_sub.font.name = "Arial Narrow"
    r_sub.font.size = Pt(15)
    r_sub.font.bold = True
    r_sub.font.color.rgb = GOLD

    add_separator(doc, color=GOLD, sz=8)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(8)
    for i, line in enumerate(metadados_lines):
        if i > 0:
            p_meta.add_run().add_break()
        r = p_meta.add_run(line)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED_C

    if alerta_lines:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        add_alert(doc, alerta_lines, kind="warn", label="Aviso")

    page_break(doc)


def add_sumario(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("SUMÁRIO")
    r.font.name = "Arial Narrow"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = NAVY
    insert_toc(doc)
    page_break(doc)
