"""
Converte Manual_Calculadora_PAR_v8.docx para o padrão visual COGER.
Preserva TODO o conteúdo, apenas reaplica formatação.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

SRC = '/root/.claude/uploads/b3b6c7da-b2ac-5cf9-a028-3e4a4676d512/6b01bfeb-Manual_Calculadora_PAR_v8.docx'
OUT = '/tmp/Manual_Calculadora_PAR_COGER.docx'
LOGO = '/tmp/logo_coger.png'

NAVY    = RGBColor(0x1E, 0x3A, 0x5F)
GOLD    = RGBColor(0xBF, 0xA0, 0x40)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
MUTED_C = RGBColor(0x70, 0x70, 0x70)
BLUE_D  = RGBColor(0x2E, 0x75, 0xB6)
GREEN   = RGBColor(0x1A, 0x6E, 0x1A)
GREEN_B = 'E8F5E9'

def rgb_hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

# ── XML helpers ───────────────────────────────────────────────────────────────
def set_shading_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color.upper())
    for old in pPr.findall(qn('w:shd')): pPr.remove(old)
    pPr.append(shd)

def set_shading_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color.upper())
    for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
    tcPr.append(shd)

def set_left_border(para, color, sz=18):
    pPr  = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None: pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),str(sz))
    left.set(qn('w:space'),'6');    left.set(qn('w:color'),rgb_hex(color))
    for old in pBdr.findall(qn('w:left')): pBdr.remove(old)
    pBdr.append(left)

def set_cell_borders(cell, color='CCCCCC'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),color)
        tcBdr.append(el)
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    tcPr.append(tcBdr)

def add_runs_from_source(para, src_para, size=11, font='Calibri', color=None):
    """Copy runs from a source paragraph preserving bold/italic."""
    runs = src_para.runs
    if not runs:
        text = src_para.text
        if text.strip():
            r = para.add_run(text)
            r.font.name = font; r.font.size = Pt(size)
            if color: r.font.color.rgb = color
        return
    for src_r in runs:
        if not src_r.text: continue
        r = para.add_run(src_r.text)
        r.font.name = font; r.font.size = Pt(size)
        r.font.bold = src_r.bold
        r.font.italic = src_r.italic
        if color: r.font.color.rgb = color

def add_runs(para, text, size=11, font='Calibri', color=None, base_bold=False, base_italic=False):
    parts = re.split(r'(\*\*|\*)', text)
    bold = base_bold; italic = base_italic
    for seg in parts:
        if seg == '**': bold   = not bold;   continue
        if seg == '*':  italic = not italic; continue
        if not seg: continue
        r = para.add_run(seg)
        r.font.name=font; r.font.size=Pt(size); r.font.bold=bold; r.font.italic=italic
        if color: r.font.color.rgb = color

# ── Style setup ───────────────────────────────────────────────────────────────
def setup_heading_styles(doc):
    def _style(name, font_name, size_pt, bold, color, space_before, space_after,
               left_border_color=None, left_border_sz=None):
        s = doc.styles[name]
        s.font.name      = font_name
        s.font.size      = Pt(size_pt)
        s.font.bold      = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(space_before)
        s.paragraph_format.space_after  = Pt(space_after)
        s.paragraph_format.keep_with_next = True
        pPr = s.element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None: pPr.remove(numPr)
        if left_border_color and left_border_sz:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None: pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
            left = OxmlElement('w:left')
            left.set(qn('w:val'),'single')
            left.set(qn('w:sz'), str(left_border_sz))
            left.set(qn('w:space'),'6')
            left.set(qn('w:color'), rgb_hex(left_border_color))
            for old in pBdr.findall(qn('w:left')): pBdr.remove(old)
            pBdr.append(left)
    _style('Heading 1', 'Arial Narrow', 20, True, NAVY, 4, 12)
    _style('Heading 2', 'Arial Narrow', 13, True, NAVY, 12, 5,
           left_border_color=GOLD, left_border_sz=18)
    _style('Heading 3', 'Calibri', 11, True, NAVY, 8, 3)

def insert_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)
    placeholder = OxmlElement('w:t')
    placeholder.text = 'Clique com o botão direito → Atualizar campo para ver o sumário com páginas.'
    run._r.append(placeholder)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)

# ── Components ────────────────────────────────────────────────────────────────
def add_chapter(doc, title, number=None):
    if number is not None:
        p_eye = doc.add_paragraph()
        p_eye.paragraph_format.space_before=Pt(4); p_eye.paragraph_format.space_after=Pt(0)
        r_eye = p_eye.add_run(f'CAPÍTULO {number}')
        r_eye.font.name='Arial Narrow'; r_eye.font.size=Pt(9)
        r_eye.font.bold=True; r_eye.font.color.rgb=GOLD
    p2 = doc.add_paragraph(style='Heading 1')
    p2.paragraph_format.space_before=Pt(2); p2.paragraph_format.space_after=Pt(12)
    r2 = p2.add_run(title)
    r2.font.name='Arial Narrow'; r2.font.size=Pt(20)
    r2.font.bold=True; r2.font.color.rgb=NAVY
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before=Pt(0); sep.paragraph_format.space_after=Pt(10)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),rgb_hex(GOLD))
    pBdr.append(bot); pPr.append(pBdr)

def add_h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(5)
    set_left_border(p, GOLD, 18)
    r = p.add_run(text)
    r.font.name='Arial Narrow'; r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=NAVY

def add_h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
    r = p.add_run(text)
    r.font.name='Calibri'; r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=NAVY

def add_body_para(doc, src_para):
    p = doc.add_paragraph()
    p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs_from_source(p, src_para)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)

def add_alert(doc, lines, kind='info', label=None):
    bg  = {'info':'EBF3FB','warn':'FFF8E1','danger':'FDE8E8','green':GREEN_B,'mono':'F4F4F4'}
    bdc = {'info':BLUE_D,'warn':GOLD,'danger':RGBColor(0xC0,0x00,0x00),'green':GREEN,'mono':MUTED_C}
    bg_c = bg.get(kind,'EBF3FB'); bd = bdc.get(kind,BLUE_D)
    if isinstance(lines, str): lines = [lines]
    first = True
    for line in lines:
        if not line.strip(): continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent=Cm(0.4); p.paragraph_format.right_indent=Cm(0.3)
        p.paragraph_format.space_before=Pt(5) if first else Pt(1)
        p.paragraph_format.space_after=Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_left_border(p, bd, 24); set_shading_para(p, bg_c)
        if first and label:
            r0 = p.add_run(f'{label}  ')
            r0.font.name='Arial Narrow'; r0.font.size=Pt(9); r0.font.bold=True; r0.font.color.rgb=bd
        add_runs(p, line, size=10)
        first = False
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def add_alert_para(doc, src_paras, kind='warn', label=None):
    """Alert box from source paragraphs (preserving runs)."""
    bg  = {'info':'EBF3FB','warn':'FFF8E1','danger':'FDE8E8','green':GREEN_B,'mono':'F4F4F4'}
    bdc = {'info':BLUE_D,'warn':GOLD,'danger':RGBColor(0xC0,0x00,0x00),'green':GREEN,'mono':MUTED_C}
    bg_c = bg.get(kind,'EBF3FB'); bd = bdc.get(kind,BLUE_D)
    if not isinstance(src_paras, list): src_paras = [src_paras]
    first = True
    for src_para in src_paras:
        if not src_para.text.strip(): continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent=Cm(0.4); p.paragraph_format.right_indent=Cm(0.3)
        p.paragraph_format.space_before=Pt(5) if first else Pt(1)
        p.paragraph_format.space_after=Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_left_border(p, bd, 24); set_shading_para(p, bg_c)
        if first and label:
            r0 = p.add_run(f'{label}  ')
            r0.font.name='Arial Narrow'; r0.font.size=Pt(9); r0.font.bold=True; r0.font.color.rgb=bd
        add_runs_from_source(p, src_para, size=10)
        first = False
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def add_part_divider(doc, text):
    """PARTE I / PARTE II divider — centered, large navy."""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(20)
    r = p.add_run(text)
    r.font.name='Arial Narrow'; r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=NAVY

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),rgb_hex(GOLD))
    pBdr.append(bot); pPr.append(pBdr)

def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def make_table(doc, headers, rows, col_widths, alt_bg='F5F8FF', header_bg=None):
    if header_bg is None: header_bg = rgb_hex(NAVY)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for j, (h, w) in enumerate(zip(headers, col_widths)):
        c = hdr.cells[j]; c.width = Cm(w)
        set_shading_cell(c, header_bg); set_cell_borders(c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        add_runs(p, h, size=9, color=WHITE, base_bold=True)
    for i, row in enumerate(rows):
        r = t.add_row()
        bg = alt_bg if i % 2 == 0 else 'FFFFFF'
        for j, (val, w) in enumerate(zip(row, col_widths)):
            c = r.cells[j]; c.width = Cm(w)
            set_shading_cell(c, bg); set_cell_borders(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            add_runs(p, str(val), size=9.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

def make_table_from_src(doc, src_table):
    """Build a COGER-styled table from a source python-docx Table object."""
    rows_data = []
    for r in src_table.rows:
        rows_data.append([c.text for c in r.cells])
    if not rows_data: return
    ncols = max(len(r) for r in rows_data)
    # Distribute column widths evenly to 16cm
    col_widths = [16.0 / ncols] * ncols
    headers = rows_data[0]
    rows = rows_data[1:]
    make_table(doc, headers, rows, col_widths)

# ── Helpers ───────────────────────────────────────────────────────────────────
def get_para_style(p):
    """Return normalized style name."""
    try:
        return p.style.name
    except Exception:
        return 'Normal'

def is_gray_fill(cell):
    """True if cell has a gray/near-gray fill (example boxes)."""
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None: return False
    shd = tcPr.find(qn('w:shd'))
    if shd is None: return False
    fill = shd.get(qn('w:fill'), '').upper()
    if not fill or fill == 'AUTO': return False
    try:
        r = int(fill[0:2], 16)
        g = int(fill[2:4], 16)
        b = int(fill[4:6], 16)
        # Gray-ish: all channels close to each other and >= 200
        return abs(r-g) < 30 and abs(g-b) < 30 and r >= 180
    except Exception:
        return False

def is_1x1_gray(tbl):
    try:
        if len(tbl.rows) == 1 and len(tbl.rows[0].cells) == 1:
            return is_gray_fill(tbl.rows[0].cells[0])
    except Exception:
        pass
    return False

RE_NUMBERED_H1 = re.compile(r'^(\d+)\.\s+(.+)')

def classify_h1(text):
    """Return ('chapter', num, title) or ('part', text) or ('section', text) or ('alert', text)."""
    t = text.strip()
    if t.startswith('⚠') or t.startswith('ATENÇÃO') or t.startswith('Atenção'):
        return ('alert', t)
    if re.match(r'^PARTE\s+(I|II|III|IV|V)\b', t, re.IGNORECASE):
        return ('part', t)
    m = RE_NUMBERED_H1.match(t)
    if m:
        num = m.group(1)
        title_rest = m.group(2).strip()
        # If mostly uppercase → chapter heading; if has many lowercase → list item
        lowers = sum(1 for c in title_rest if c.islower())
        if lowers > 5:
            return ('numbered_item', num, title_rest)
        else:
            return ('chapter', num, title_rest)
    return ('section', t)

# ── COVER DATA ────────────────────────────────────────────────────────────────
COVER_TITLE = 'MANUAL TEÓRICO-PRÁTICO'
COVER_SUBTITLE = 'Calculadora de Sanções — PAR'
COVER_REFS = 'Lei nº 8.112/1990  ·  Lei nº 9.784/1999  ·  Lei nº 8.429/1992'
COVER_VERSION = 'Versão 8'
COVER_AUTHORS = 'Corregedoria da Receita Federal do Brasil'

# ── BUILD ─────────────────────────────────────────────────────────────────────
def build():
    src = Document(SRC)
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_height=Cm(29.7); sec.page_width=Cm(21.0)
    sec.top_margin=Cm(2.2); sec.bottom_margin=Cm(2.2)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    sec.header.is_linked_to_previous=False
    sec.footer.is_linked_to_previous=False
    doc.styles['Normal'].font.name='Calibri'
    doc.styles['Normal'].font.size=Pt(11)
    doc.styles['Normal'].paragraph_format.space_after=Pt(5)
    setup_heading_styles(doc)

    # ── CAPA ─────────────────────────────────────────────────────────────────
    try:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(20)
        p_logo.paragraph_format.space_after  = Pt(14)
        p_logo.add_run().add_picture(LOGO, width=Cm(7))
    except Exception as e:
        print(f'[WARN] Logo: {e}')

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(4)
    r = p_title.add_run(COVER_TITLE)
    r.font.name='Arial Narrow'; r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(2)
    p_sub.paragraph_format.space_after  = Pt(8)
    r2 = p_sub.add_run(COVER_SUBTITLE)
    r2.font.name='Arial Narrow'; r2.font.size=Pt(16); r2.font.bold=True; r2.font.color.rgb=GOLD

    add_separator(doc)

    for txt, sz, color in [
        (COVER_REFS, 10, MUTED_C),
        (COVER_VERSION, 11, NAVY),
        (COVER_AUTHORS, 10, MUTED_C),
    ]:
        p_x = doc.add_paragraph()
        p_x.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_x.paragraph_format.space_after = Pt(3)
        r_x = p_x.add_run(txt)
        r_x.font.name='Calibri'; r_x.font.size=Pt(sz); r_x.font.color.rgb=color

    page_break(doc)

    # ── SUMÁRIO ───────────────────────────────────────────────────────────────
    p_s = doc.add_paragraph()
    p_s.paragraph_format.space_before = Pt(0)
    p_s.paragraph_format.space_after  = Pt(14)
    r_s = p_s.add_run('SUMÁRIO')
    r_s.font.name='Arial Narrow'; r_s.font.size=Pt(18); r_s.font.bold=True; r_s.font.color.rgb=NAVY

    insert_toc(doc)
    page_break(doc)

    # ── BODY ─────────────────────────────────────────────────────────────────
    # Iterate body XML in document order (preserves paragraph/table interleave)
    body = src.element.body
    # Build lookup: xml element → paragraph or table object
    para_map = {p._p: p for p in src.paragraphs}
    tbl_map  = {t._tbl: t for t in src.tables}

    # Track state for chapter numbering override
    part_counters = {}   # just pass through text as-is

    SKIP_STYLES = {'toc 1', 'toc 2', 'toc 3', 'TOC 1', 'TOC 2', 'TOC 3'}
    in_toc = False

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = para_map.get(child)
            if para is None: continue
            style_name = get_para_style(para)
            text = para.text.strip()

            # Skip empty paragraphs (except as spacers — we add our own spacing)
            if not text: continue

            # Skip TOC entries
            if style_name in SKIP_STYLES: continue
            if style_name == 'Heading 1' and text == 'SUMÁRIO': continue

            # Heading 1 classification
            if style_name == 'Heading 1':
                result = classify_h1(text)
                kind = result[0]
                if kind == 'alert':
                    body_text = result[1].lstrip('⚠').strip()
                    add_alert(doc, [body_text], kind='warn', label='ATENÇÃO')
                elif kind == 'part':
                    page_break(doc)
                    add_part_divider(doc, result[1])
                elif kind == 'chapter':
                    page_break(doc)
                    add_chapter(doc, result[1], result[2])
                elif kind == 'numbered_item':
                    # numbered list item written as Heading 1
                    p_n = doc.add_paragraph()
                    p_n.paragraph_format.left_indent = Cm(0.8)
                    p_n.paragraph_format.space_after = Pt(5)
                    p_n.paragraph_format.line_spacing = Pt(13.2)
                    p_n.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    r_n = p_n.add_run(f'{result[1]}.  ')
                    r_n.font.name='Calibri'; r_n.font.size=Pt(11)
                    r_n.font.bold=True; r_n.font.color.rgb=NAVY
                    r_txt = p_n.add_run(result[2])
                    r_txt.font.name='Calibri'; r_txt.font.size=Pt(11)
                else:  # section
                    page_break(doc)
                    add_chapter(doc, text)

            elif style_name == 'Heading 2':
                add_h2(doc, text)

            elif style_name in ('Heading 3', 'Heading 4'):
                add_h3(doc, text)

            elif style_name == 'Intense Quote':
                add_alert_para(doc, [para], kind='warn', label='ESTADO DA CONTROVÉRSIA')

            else:  # Normal and anything else
                # Normal starting with ⚠
                if text.startswith('⚠') or text.startswith('Atenção:') or text.startswith('ATENÇÃO:'):
                    body_text = re.sub(r'^[⚠\s]*(Atenção|ATENÇÃO)\s*:\s*', '', text).strip()
                    if not body_text: body_text = text.lstrip('⚠').strip()
                    add_alert(doc, [body_text], kind='warn', label='ATENÇÃO')
                else:
                    add_body_para(doc, para)

        elif tag == 'tbl':
            tbl = tbl_map.get(child)
            if tbl is None: continue

            if is_1x1_gray(tbl):
                # Example callout box
                cell = tbl.rows[0].cells[0]
                cell_paras = cell.paragraphs
                # Find label from first non-empty paragraph that has bold runs
                label = None
                content_paras = []
                for i, cp in enumerate(cell_paras):
                    ct = cp.text.strip()
                    if not ct: continue
                    if label is None:
                        # Check if first para is bold (label)
                        is_bold = any(r.bold for r in cp.runs if r.text.strip())
                        if is_bold:
                            label = ct
                        else:
                            content_paras.append(cp)
                    else:
                        content_paras.append(cp)

                if not content_paras and label:
                    # Only one paragraph — it's both label and content
                    add_alert(doc, [label], kind='green', label='Exemplo')
                else:
                    lines = [cp.text for cp in content_paras if cp.text.strip()]
                    add_alert(doc, lines, kind='green', label=label or 'Exemplo')
            else:
                # Multi-column data table
                make_table_from_src(doc, tbl)

    doc.save(OUT)
    print(f'Saved: {OUT}')
    # Count output elements
    out_doc = Document(OUT)
    print(f'  Paragraphs: {len(out_doc.paragraphs)}, Tables: {len(out_doc.tables)}')

if __name__ == '__main__':
    build()
