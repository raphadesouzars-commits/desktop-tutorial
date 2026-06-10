from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Margens ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ── Estilos de parágrafo ─────────────────────────────────────────────────────
def set_font(run, size, bold=False, italic=False, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 16, bold=True, color=(0, 51, 102))
    return p

def add_tela(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), '003366')
    p._p.get_or_add_pPr().append(shading_elm)
    run = p.add_run(f"  TELA {num}: {title}  ")
    set_font(run, 12, bold=True, color=(255, 255, 255))
    return p

def add_modulo(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), 'D0E4F7')
    p._p.get_or_add_pPr().append(shading_elm)
    run = p.add_run(f"  {text}  ")
    set_font(run, 10, bold=True, color=(0, 51, 102))
    return p

def add_campo(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(f"▶  {text}")
    set_font(run, 9.5, bold=True, color=(0, 80, 160))
    return p

def add_opcao(doc, text, pts, nivel=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0 * nivel)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    bullet = "●" if nivel == 1 else "○" if nivel == 2 else "–"
    run_b = p.add_run(f"  {bullet}  ")
    set_font(run_b, 9)
    run_t = p.add_run(text)
    set_font(run_t, 9)
    if pts is not None:
        run_p = p.add_run(f"  →  {pts}")
        set_font(run_p, 9, bold=True, color=(180, 0, 0))
    return p

def add_obs(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(f"ℹ  {text}")
    set_font(run, 8.5, italic=True, color=(100, 100, 100))
    return p

def add_sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 110)
    set_font(run, 7, color=(180, 180, 180))

# ═══════════════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════════════
add_title(doc, "CALCULADORA DE DOSIMETRIA TAC")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Receita Federal do Brasil — Corregedoria")
set_font(r, 11, color=(80, 80, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Árvore Completa de Opções e Pontuações por Tela")
set_font(r2, 11, bold=True, color=(0, 51, 102))

doc.add_paragraph()
add_obs(doc, "Este documento lista todas as telas da calculadora, com cada campo, opção possível e pontuação correspondente.")
add_obs(doc, "Pontuação máxima total: 120 pts  (Natureza 30 + Gravidade 30 + Dano 30 + Circunstâncias ±15 + Antecedentes ±15)")
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 0 — IDENTIFICAÇÃO
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 0, "IDENTIFICAÇÃO")
add_obs(doc, "Campos informativos — sem pontuação.")
campos_id = [
    "Número do Processo",
    "Data do Ilícito",
    "Nome do Servidor",
    "Matrícula SIAPE",
    "Unidade / Delegacia",
    "Cargo / Função",
    "Confirmação de Escopo (checkbox obrigatório)",
]
for c in campos_id:
    add_campo(doc, c)
    add_opcao(doc, "Preenchimento livre (texto/data)", None)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 1 — CLASSIFICAÇÃO DA INFRAÇÃO
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 1, "CLASSIFICAÇÃO DA INFRAÇÃO")
add_obs(doc, "Sem pontuação direta — define o regime e o enquadramento legal aplicável.")

add_campo(doc, "Regime da Infração")
regimes = [
    ("A — Regime Geral",        "Art. 116/117 Lei 8.112/90 → Advertência ou Suspensão 1–90 dias"),
    ("B — Regime RE1",          "Art. 117 XVII–XVIII Lei 8.112/90 → Suspensão 1–90 dias"),
    ("C — Regime RE2",          "Art. 130, §1º Lei 8.112/90 → Suspensão 1–15 dias"),
    ("D — Regime RE3",          "Art. 32 Lei 12.527/2011 (LAI) → Suspensão 1–90 dias"),
    ("E — Concurso de Infrações","Múltiplas infrações — fluxo especial de cálculo"),
]
for label, desc in regimes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run_b = p.add_run(f"  ●  {label}")
    set_font(run_b, 9, bold=True)
    run_d = p.add_run(f"  —  {desc}")
    set_font(run_d, 9, italic=True, color=(80, 80, 80))

add_obs(doc, "Após escolher o regime, abre modal para seleção do enquadramento legal (inciso/artigo específico).")
add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 2 — NATUREZA DA INFRAÇÃO  (Teto: 30 pts)
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 2, "NATUREZA DA INFRAÇÃO")
add_modulo(doc, "Elemento 1 · Módulo 1.1 | TETO: 30 pontos")

# SF1.1
add_campo(doc, "SF1.1 — Elemento Subjetivo Principal  (0–20 pts)")

add_opcao(doc, "CULPA LEVE  [faixa: 1–6 pts]", None)
add_opcao(doc, "Descuido pontual, sem repercussão relevante", "2 pts", nivel=2)
add_opcao(doc, "Negligência perceptível, com alguma repercussão", "4 pts", nivel=2)
add_opcao(doc, "Negligência acentuada, aquém do erro grosseiro", "6 pts", nivel=2)

add_opcao(doc, "CULPA GRAVE / ERRO GROSSEIRO  [faixa: 7–14 pts]", None)
add_opcao(doc, "Evidente para a maioria, com complexidade técnica", "8 pts", nivel=2)
add_opcao(doc, "Claramente evidente, sem complexidade técnica", "11 pts", nivel=2)
add_opcao(doc, "Evidente e extremamente grave", "14 pts", nivel=2)

add_opcao(doc, "DOLO  [faixa: 15–20 pts]", None)
add_opcao(doc, "Dolo eventual (assumiu o risco)", "15 pts", nivel=2)
add_opcao(doc, "Dolo direto com resultado de menor gravidade", "17 pts", nivel=2)
add_opcao(doc, "Dolo direto com planejamento / alta gravidade", "20 pts", nivel=2)

# SF1.2
add_campo(doc, "SF1.2 — Consciência da Ilicitude  (0–6 pts)")

add_opcao(doc, "SEM NORMA ACESSÍVEL  [faixa: 0–2 pts]", None)
add_opcao(doc, "Ausência total de orientação formal", "0 pts", nivel=2)
add_opcao(doc, "Havia referências informais", "2 pts", nivel=2)

add_opcao(doc, "HAVIA NORMA GERAL  [faixa: 3–4 pts]", None)
add_opcao(doc, "Norma técnica de difícil compreensão", "3 pts", nivel=2)
add_opcao(doc, "Norma clara e amplamente divulgada", "4 pts", nivel=2)

add_opcao(doc, "NORMA + TREINAMENTO  [faixa: 5–6 pts]", None)
add_opcao(doc, "Treinamento geral na área", "5 pts", nivel=2)
add_opcao(doc, "Treinamento específico documentado", "6 pts", nivel=2)

# SF1.3
add_campo(doc, "SF1.3 — Modalidade do Dolo  (0–4 pts)  ⚠ Exibido apenas se SF1.1 ≥ 15 pts")

add_opcao(doc, "DOLO EVENTUAL  [faixa: 1–2 pts]", None)
add_opcao(doc, "Risco previsível, mas não plenamente dimensionado", "1 pt", nivel=2)
add_opcao(doc, "Risco claramente previsível e conscientemente desconsiderado", "2 pts", nivel=2)

add_opcao(doc, "DOLO DIRETO  [faixa: 3–4 pts]", None)
add_opcao(doc, "Resultado buscado, sem evidência de planejamento", "3 pts", nivel=2)
add_opcao(doc, "Resultado buscado com planejamento verificável", "4 pts", nivel=2)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 3 — GRAVIDADE DA INFRAÇÃO  (Teto: 30 pts)
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 3, "GRAVIDADE DA INFRAÇÃO")
add_modulo(doc, "Elemento 2 · Módulo 1.2 | TETO: 30 pontos")

# SF2.1
add_campo(doc, "SF2.1 — Hierarquia do Bem Jurídico Protegido  (1–10 pts)")

add_opcao(doc, "BAIXA PROTEÇÃO  [faixa: 1–3 pts]", None)
add_opcao(doc, "Não tem reflexo na imagem da unidade", "1 pt", nivel=2)
add_opcao(doc, "Algum reflexo na imagem da unidade", "3 pts", nivel=2)

add_opcao(doc, "MÉDIA PROTEÇÃO  [faixa: 4–7 pts]", None)
add_opcao(doc, "Procedimental sem impacto direto", "4 pts", nivel=2)
add_opcao(doc, "Integra sistema de controle interno", "6 pts", nivel=2)
add_opcao(doc, "Integra cadeia de responsabilização formal", "7 pts", nivel=2)

add_opcao(doc, "ALTA PROTEÇÃO  [faixa: 8–10 pts]", None)
add_opcao(doc, "Regularidade formal fiscal / aduaneira", "8 pts", nivel=2)
add_opcao(doc, "Integridade competitiva de licitações", "9 pts", nivel=2)
add_opcao(doc, "Sigilo fiscal / segurança de sistemas", "10 pts", nivel=2)

# SF2.2
add_campo(doc, "SF2.2 — Caráter Temporal da Conduta  (1–8 pts)")

add_opcao(doc, "ISOLADA  [faixa: 1–3 pts]", None)
add_opcao(doc, "Episódio único, interrompido voluntariamente", "1 pt", nivel=2)
add_opcao(doc, "Episódio único, sem voluntariedade", "2 pts", nivel=2)
add_opcao(doc, "Episódio único, mas se repetiria", "3 pts", nivel=2)

add_opcao(doc, "REITERADA  [faixa: 4–6 pts]", None)
add_opcao(doc, "2–3 episódios em até 3 meses", "4 pts", nivel=2)
add_opcao(doc, "4+ episódios ou período >3 meses", "6 pts", nivel=2)

add_opcao(doc, "CONTINUADA  [faixa: 6–8 pts]", None)
add_opcao(doc, "Até 6 meses de continuidade", "6 pts", nivel=2)
add_opcao(doc, "Mais de 6 meses de continuidade", "8 pts", nivel=2)

# SF2.3
add_campo(doc, "SF2.3 — Alcance da Conduta  (1–7 pts)")

add_opcao(doc, "RESTRITO  [faixa: 1–2 pts]", None)
add_opcao(doc, "Exclusivamente individual", "1 pt", nivel=2)
add_opcao(doc, "Algum impacto na unidade imediata", "2 pts", nivel=2)

add_opcao(doc, "AFETOU O ÓRGÃO  [faixa: 3–5 pts]", None)
add_opcao(doc, "Afetou outras unidades pontualmente", "3 pts", nivel=2)
add_opcao(doc, "Afetou área ou divisão com abrangência", "5 pts", nivel=2)

add_opcao(doc, "AFETOU EXTERNOS  [faixa: 5–7 pts]", None)
add_opcao(doc, "Afetou contribuinte identificável", "5 pts", nivel=2)
add_opcao(doc, "Afetou grupo ou gerou repercussão externa", "7 pts", nivel=2)

# SF2.4
add_campo(doc, "SF2.4 — Concurso de Dispositivos Violados  (0–5 pts)")
add_opcao(doc, "Um único dispositivo", "0 pts")
add_opcao(doc, "Dois dispositivos", "2 pts")
add_opcao(doc, "Três dispositivos", "3 pts")
add_opcao(doc, "Quatro dispositivos", "4 pts")
add_opcao(doc, "Cinco ou mais dispositivos", "5 pts")

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 4 — DANO MATERIAL  (Teto individual: 30 pts | Teto conjunto c/ imaterial: 30 pts)
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 4, "DANO MATERIAL")
add_modulo(doc, "Subelemento 3A · Módulo 1.3A | TETO individual: 30 pts | Teto conjunto c/ dano imaterial: 30 pts")

add_campo(doc, "SF3A.1 — Existência e Extensão do Dano Patrimonial")

add_opcao(doc, "SEM DANO MATERIAL", "0 pts")

add_opcao(doc, "DANO POTENCIAL  [faixa: 3–8 pts]", None)
add_opcao(doc, "Baixo impacto potencial (<5% do valor em risco)", "3 pts", nivel=2)
add_opcao(doc, "Médio impacto potencial (5–25%)", "6 pts", nivel=2)
add_opcao(doc, "Alto impacto potencial (>25%)", "8 pts", nivel=2)

add_opcao(doc, "DANO EFETIVO — IMPACTO BAIXO  [faixa: 9–15 pts]", None)
add_opcao(doc, "Integralmente ressarcido", "9 pts", nivel=2)
add_opcao(doc, "Parcialmente reparado", "12 pts", nivel=2)
add_opcao(doc, "Não ressarcido", "15 pts", nivel=2)

add_opcao(doc, "DANO EFETIVO — IMPACTO MÉDIO  [faixa: 16–22 pts]", None)
add_opcao(doc, "Integralmente ressarcido", "16 pts", nivel=2)
add_opcao(doc, "Parcialmente reparado", "19 pts", nivel=2)
add_opcao(doc, "Não ressarcido", "22 pts", nivel=2)

add_opcao(doc, "DANO EFETIVO — IMPACTO ALTO  [faixa: 23–30 pts]", None)
add_opcao(doc, "Integralmente ressarcido", "23 pts", nivel=2)
add_opcao(doc, "Parcialmente reparado", "27 pts", nivel=2)
add_opcao(doc, "Não ressarcido / irreversível", "30 pts", nivel=2)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 5 — DANO IMATERIAL  (Teto individual: 30 pts | Teto conjunto: 30 pts)
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 5, "DANO IMATERIAL")
add_modulo(doc, "Subelemento 3B · Módulo 1.3B | TETO individual: 30 pts | Teto conjunto c/ dano material: 30 pts")

# SF3B.1
add_campo(doc, "SF3B.1 — Impacto na Imagem / Credibilidade da RFB  (0–20 pts)")
add_opcao(doc, "Sem impacto", "0 pts")
add_opcao(doc, "REPERCUSSÃO LIMITADA  [faixa: 4–12 pts]", None)
add_opcao(doc, "Restrito à unidade", "4 pts", nivel=2)
add_opcao(doc, "Repercussão regional", "8 pts", nivel=2)
add_opcao(doc, "Repercussão institucional (RFB como um todo)", "12 pts", nivel=2)
add_opcao(doc, "REPERCUSSÃO AMPLA  [faixa: 13–20 pts]", None)
add_opcao(doc, "Potencial concreto de repercussão, sem efetivação", "13 pts", nivel=2)
add_opcao(doc, "Repercussão efetiva verificada", "20 pts", nivel=2)

# SF3B.2
add_campo(doc, "SF3B.2 — Impacto na Regularidade do Serviço Tributário / Aduaneiro  (0–10 pts)")
add_opcao(doc, "Sem impacto", "0 pts")
add_opcao(doc, "TURBAÇÃO PONTUAL  [faixa: 3–6 pts]", None)
add_opcao(doc, "Turbação, mas resultado preservado", "3 pts", nivel=2)
add_opcao(doc, "Resultado parcialmente comprometido", "5 pts", nivel=2)
add_opcao(doc, "Resultado comprometido, com necessidade de revisão", "6 pts", nivel=2)
add_opcao(doc, "COMPROMETIMENTO EFETIVO  [faixa: 7–10 pts]", None)
add_opcao(doc, "Sem impacto na arrecadação", "7 pts", nivel=2)
add_opcao(doc, "Com impacto direto na arrecadação", "10 pts", nivel=2)

# SF3B.3
add_campo(doc, "SF3B.3 — Impacto sobre Contribuintes / Usuários / Terceiros  (0–10 pts)")
add_opcao(doc, "Sem impacto", "0 pts")
add_opcao(doc, "INDIVIDUAL  [faixa: 2–5 pts]", None)
add_opcao(doc, "Perturbação sem prejuízo concreto", "2 pts", nivel=2)
add_opcao(doc, "Prejuízo concreto verificado", "5 pts", nivel=2)
add_opcao(doc, "DIFUSO  [faixa: 6–10 pts]", None)
add_opcao(doc, "Grupo identificável afetado", "6 pts", nivel=2)
add_opcao(doc, "Impacto difuso sobre coletividade", "10 pts", nivel=2)

# SF3B.4a
add_campo(doc, "SF3B.4a — Comprometimento do Sigilo Fiscal  (0–20 pts)")
add_opcao(doc, "Sem comprometimento", "0 pts")
add_opcao(doc, "RISCO CONCRETO  [faixa: 4–8 pts]", None)
add_opcao(doc, "Acesso indevido a dados de contribuinte individual", "4 pts", nivel=2)
add_opcao(doc, "Acesso indevido a dados de grupo de contribuintes", "8 pts", nivel=2)
add_opcao(doc, "EXPOSIÇÃO EFETIVA  [faixa: 9–20 pts]", None)
add_opcao(doc, "Exposição de dados de contribuinte individual", "9 pts", nivel=2)
add_opcao(doc, "Exposição de dados de grupo de contribuintes", "14 pts", nivel=2)
add_opcao(doc, "Exposição ampla / uso ilícito das informações", "20 pts", nivel=2)

# SF3B.4b
add_campo(doc, "SF3B.4b — Comprometimento do Processo Tributário / Aduaneiro  (0–20 pts)")
add_opcao(doc, "Sem comprometimento", "0 pts")
add_opcao(doc, "RISCO À ISONOMIA  [faixa: 3–8 pts]", None)
add_opcao(doc, "Risco hipotético à isonomia", "3 pts", nivel=2)
add_opcao(doc, "Risco iminente à isonomia", "8 pts", nivel=2)
add_opcao(doc, "COMPROMETIMENTO DIRETO  [faixa: 9–20 pts]", None)
add_opcao(doc, "Sem vantagem indevida para o infrator", "9 pts", nivel=2)
add_opcao(doc, "Com vantagem indevida para o infrator", "14 pts", nivel=2)
add_opcao(doc, "Com prejuízo direto à arrecadação", "20 pts", nivel=2)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 6 — CIRCUNSTÂNCIAS AGRAVANTES E ATENUANTES  (Modificador: ±15 pts)
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 6, "CIRCUNSTÂNCIAS AGRAVANTES E ATENUANTES")
add_modulo(doc, "Elemento 4 · Módulo 1.4 | Modificador: até +15 pts (agravantes) / até −15 pts (atenuantes)")
add_obs(doc, "Cada circunstância possui uma faixa de pontos e um valor padrão (default). O usuário pode ajustar dentro da faixa.")

add_campo(doc, "AGRAVANTES  (somam pontos — teto acumulado: +15 pts)")
agravantes = [
    ("AG1", "Capacitação específica documentada para a função", "3–5 pts", "padrão: 4"),
    ("AG2", "Longa experiência no cargo (>3 anos)", "3–4 pts", "padrão: 3"),
    ("AG3", "Exercício de cargo de confiança / chefia", "3–5 pts", "padrão: 4"),
    ("AG4", "Interesse fútil ou inescusável na infração", "2–4 pts", "padrão: 3"),
    ("AG5", "Infraestrutura favorável disponível não utilizada corretamente", "2–3 pts", "padrão: 2"),
    ("AG6", "Conduta em detrimento de pessoa vulnerável", "2–4 pts", "padrão: 3"),
]
for cod, desc, faixa, pad in agravantes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"  ●  [{cod}]  ")
    set_font(r1, 9, bold=True, color=(180, 0, 0))
    r2 = p.add_run(desc)
    set_font(r2, 9)
    r3 = p.add_run(f"   Faixa: {faixa}  |  {pad}")
    set_font(r3, 9, bold=True, color=(180, 0, 0))

add_campo(doc, "ATENUANTES  (reduzem pontos — teto acumulado: −15 pts)")
atenuantes = [
    ("AT1", "Ausência de capacitação específica documentada", "3–5 pts", "padrão: 4"),
    ("AT2", "Pouco tempo de serviço no cargo (<2 anos)", "3–4 pts", "padrão: 3"),
    ("AT3", "Sobrecarga de trabalho documentada", "2–4 pts", "padrão: 3"),
    ("AT4", "Condições precárias de infraestrutura comprovadas", "2–3 pts", "padrão: 2"),
    ("AT5", "Confissão espontânea da infração", "2–3 pts", "padrão: 2"),
    ("AT6", "Reparação voluntária do dano antes da instauração", "2–4 pts", "padrão: 3"),
    ("AT7", "Problemas pessoais documentados que influenciaram a conduta", "2–3 pts", "padrão: 2"),
    ("AT8", "Contexto de urgência / emergência que motivou a conduta", "2–4 pts", "padrão: 3"),
]
for cod, desc, faixa, pad in atenuantes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"  ●  [{cod}]  ")
    set_font(r1, 9, bold=True, color=(0, 120, 0))
    r2 = p.add_run(desc)
    set_font(r2, 9)
    r3 = p.add_run(f"   Faixa: −{faixa}  |  {pad}")
    set_font(r3, 9, bold=True, color=(0, 120, 0))

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 7 — ANTECEDENTES FUNCIONAIS  (Modificador: ±15 pts)
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 7, "ANTECEDENTES FUNCIONAIS")
add_modulo(doc, "Elemento 5 · Módulo 1.5 | Modificador: até −15 pts (bons) / até +15 pts (maus)")

add_campo(doc, "BONS ANTECEDENTES  (reduzem pontos — teto acumulado: −15 pts)")
bons = [
    ("BA1", "Elogios simples registrados nos assentamentos", "2–4 pts", "padrão: 3"),
    ("BA2", "Menção honrosa ou prêmio formal", "4–7 pts", "padrão: 5"),
    ("BA3", "Registro formal de serviço relevante / destaque", "5–10 pts", "padrão: 7"),
    ("BA4", "Longa carreira sem ocorrências (>10 anos)", "3–6 pts", "padrão: 4"),
]
for cod, desc, faixa, pad in bons:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"  ●  [{cod}]  ")
    set_font(r1, 9, bold=True, color=(0, 120, 0))
    r2 = p.add_run(desc)
    set_font(r2, 9)
    r3 = p.add_run(f"   Faixa: −{faixa}  |  {pad}")
    set_font(r3, 9, bold=True, color=(0, 120, 0))

add_campo(doc, "MAUS ANTECEDENTES  (aumentam pontos — teto acumulado: +15 pts)")
maus = [
    ("MA1", "Atrasos / faltas não justificadas registrados", "2–4 pts", "padrão: 3"),
    ("MA2", "Descumprimento de TAC anterior", "4–7 pts", "padrão: 5"),
    ("MA3", "Sanção disciplinar anterior ainda não cancelada", "4–10 pts", "padrão: 6"),
    ("MA4", "Outros registros desabonadores nos assentamentos", "2–4 pts", "padrão: 3"),
]
for cod, desc, faixa, pad in maus:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"  ●  [{cod}]  ")
    set_font(r1, 9, bold=True, color=(180, 0, 0))
    r2 = p.add_run(desc)
    set_font(r2, 9)
    r3 = p.add_run(f"   Faixa: +{faixa}  |  {pad}")
    set_font(r3, 9, bold=True, color=(180, 0, 0))

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 8 — ELEGIBILIDADE DO TAC
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 8, "DADOS PARA ELEGIBILIDADE DO TAC")
add_modulo(doc, "Módulo 5 · Etapa 5.0 — Verificação de Admissibilidade")
add_obs(doc, "Sem pontuação — campos determinam se o TAC é admissível.")

add_campo(doc, "5.0.A — Vínculo do Servidor")
add_opcao(doc, "Efetivo (cargo permanente)", None)
add_opcao(doc, "Não efetivo (cargo em comissão, contrato etc.)", None)

add_campo(doc, "5.0.B — Registro Vigente de Penalidade Disciplinar")
add_opcao(doc, "Sim — impede ou condiciona elegibilidade", None)
add_opcao(doc, "Não", None)

add_campo(doc, "5.0.C — TAC Celebrado nos Últimos 2 Anos")
add_opcao(doc, "Não", None)
add_opcao(doc, "Sim  →  subcampo: infração anterior ao TAC vigente?", None)
add_opcao(doc, "    Sim (infração anterior ao TAC)", None, nivel=2)
add_opcao(doc, "    Não (infração posterior ao TAC)", None, nivel=2)

add_campo(doc, "5.0.D — Dano e Ressarcimento")
add_opcao(doc, "Não houve dano", None)
add_opcao(doc, "Houve dano  →  subcampo: ressarcimento?", None)
add_opcao(doc, "    Sim — dano ressarcido", None, nivel=2)
add_opcao(doc, "    Não — dano não ressarcido (pode inviabilizar TAC)", None, nivel=2)

add_campo(doc, "5.0.E — Dados para Geração do Instrumento (preenchimento livre)")
dados_e = [
    "Nome da Autoridade Compromitente",
    "Cargo / Função da Autoridade",
    "Nome do Chefe Imediato",
    "Cargo do Chefe Imediato",
    "Descrição dos Fatos (narrativa livre)",
    "Cidade / UF de Celebração",
    "Data de Celebração",
]
for d in dados_e:
    add_opcao(doc, d, None)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 9 — RESULTADO DA DOSIMETRIA
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 9, "RESULTADO DA DOSIMETRIA")
add_modulo(doc, "Módulos 2 + 3 + 5 — Cálculo Final e Enquadramento da Penalidade")

add_campo(doc, "Reincidência — 2ª Fase da Dosimetria")
add_opcao(doc, "Não há registro vigente de advertência (<3 anos) ou suspensão (<5 anos)", "Sem acréscimo")
add_opcao(doc, "Há registro vigente de advertência ou suspensão", "Agrava penalidade conforme regras de reincidência")

add_campo(doc, "Saídas calculadas automaticamente (sem escolha do usuário)")
saidas = [
    "Pontuação Total (soma de todos os módulos)",
    "Penalidade cabível: Advertência  ou  Suspensão X dias",
    "Viabilidade do TAC: Admissível / Não admissível / Admissível com condições",
    "Geração da Minuta do TAC (botão)",
]
for s in saidas:
    add_opcao(doc, s, None)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 10 — REVISÃO DAS OBRIGAÇÕES DO TAC
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 10, "REVISÃO DAS OBRIGAÇÕES DO TAC")
add_modulo(doc, "Módulo 5 · Etapa 5.3 — Catálogo de Obrigações (9 opções)")
add_obs(doc, "As obrigações sugeridas variam conforme a faixa da penalidade (leve / leve_media / limite).")
add_obs(doc, "Para cada obrigação selecionada: preencher Descrição, Prazo, Forma de Comprovação e Responsável pela Fiscalização.")

obrigacoes = [
    ("OB-01", "Reparação do Dano",                   "leve, leve_media, limite", "6 meses",  "Art. 68, §1º, I — Portaria CGU 27/2022"),
    ("OB-02", "Retratação Formal",                    "leve, leve_media, limite", "1 mês",    "Art. 68, §1º, II"),
    ("OB-03", "Capacitação Específica",               "leve, leve_media, limite", "6 meses",  "Art. 68, §1º, III"),
    ("OB-04", "Curso de Ética",                       "leve, leve_media",         "6 meses",  "Art. 68, §1º, III"),
    ("OB-05", "Acordo de Gestão de Horário",          "leve",                     "12 meses", "Art. 68, §1º, IV"),
    ("OB-06", "Metas de Desempenho",                  "leve_media, limite",       "12 meses", "Art. 68, §1º, V"),
    ("OB-07", "Monitoramento de Acessos a Sistemas",  "leve_media, limite",       "18 meses", "Art. 68, §1º, VI"),
    ("OB-08", "Relatório Periódico de Atividades",    "limite",                   "24 meses", "Art. 68, §1º, VI"),
    ("OB-09", "Comunicação à Corregedoria",           "limite",                   "Contínuo", "Art. 68, §1º, VI"),
]
for cod, nome, faixas, prazo, base in obrigacoes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"  ●  [{cod}]  {nome}")
    set_font(r1, 9, bold=True)
    r2 = p.add_run(f"   |  Faixas: {faixas}  |  Prazo padrão: {prazo}  |  Base: {base}")
    set_font(r2, 8.5, italic=True, color=(80, 80, 80))

add_campo(doc, "Para cada obrigação selecionada — subcampos (preenchimento livre):")
subcampos_ob = [
    "Descrição da Obrigação Efetiva",
    "Prazo Individual (1–24 meses)",
    "Forma de Comprovação",
    "Responsável pela Fiscalização",
]
for sc in subcampos_ob:
    add_opcao(doc, sc, None)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA 11 — RELATÓRIO FINAL
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, 11, "RELATÓRIO FINAL")
add_modulo(doc, "Geração automática — resumo completo da dosimetria")
add_obs(doc, "Tela de visualização e exportação. Sem campos de pontuação adicionais.")

secoes = [
    "1. Identificação (Processo, Servidor, Data, Unidade, Cargo)",
    "2. Escopo e Enquadramento Legal",
    "3. Dosimetria — 1ª Fase (detalhamento por módulo e subfator)",
    "4. Cálculo da Penalidade Bruta",
    "5. 2ª Fase — Reincidência",
    "6. Penalidade Final",
    "7. Análise de Admissibilidade do TAC (3 camadas de verificação)",
    "8. Consulta LINDB — opção de conversão em multa",
    "9. Análise de Erro Escusável (Nota Técnica Coger/RFB nº 12/2021) — opcional",
]
for s in secoes:
    add_opcao(doc, s, None)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TELA ESPECIAL — CONCURSO DE INFRAÇÕES
# ═══════════════════════════════════════════════════════════════════════════════
add_tela(doc, "E", "CONCURSO DE INFRAÇÕES  (Regime E)")
add_modulo(doc, "Fluxo especial — exibido quando o usuário seleciona Regime E na Tela 1")
add_obs(doc, "Permite cadastrar múltiplas infrações (mínimo 2) e calcular penalidade conjunta.")

add_campo(doc, "Para cada infração cadastrada:")
campos_conc = [
    "Regime da infração individual (A, B, C ou D)",
    "Pontuação calculada da infração individual",
    "Reincidência nesta infração (Sim / Não)",
]
for c in campos_conc:
    add_opcao(doc, c, None)

add_campo(doc, "Regras de cálculo do concurso (automático)")
add_opcao(doc, "Advertência + Advertência → Suspensão", None)
add_opcao(doc, "Advertência + Suspensão → Suspensão (soma de dias)", None)
add_opcao(doc, "Suspensão + Suspensão → Suspensão (soma de dias)", None)
add_opcao(doc, "Teto absoluto: 90 dias de suspensão", None)

add_sep(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# TABELA RESUMO DE PONTUAÇÕES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_title(doc, "TABELA RESUMO — PONTUAÇÕES MÁXIMAS POR MÓDULO")

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, txt in enumerate(["Módulo / Elemento", "Subfatores", "Mín.", "Máx."]):
    hdr[i].text = txt
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True

rows_data = [
    ("Módulo 1.1 — Natureza da Infração",   "SF1.1 + SF1.2 + SF1.3", "0",  "30"),
    ("Módulo 1.2 — Gravidade da Infração",  "SF2.1 + SF2.2 + SF2.3 + SF2.4", "4", "30"),
    ("Módulo 1.3A — Dano Material",         "SF3A.1", "0", "30"),
    ("Módulo 1.3B — Dano Imaterial",        "SF3B.1 + SF3B.2 + SF3B.3 + SF3B.4a/b", "0", "30"),
    ("  Teto conjunto 1.3A + 1.3B",         "—", "—", "30"),
    ("Módulo 1.4 — Circunstâncias",         "AG1–AG6 / AT1–AT8", "−15", "+15"),
    ("Módulo 1.5 — Antecedentes",           "BA1–BA4 / MA1–MA4", "−15", "+15"),
    ("TOTAL MÁXIMO POSSÍVEL",               "—", "—", "120"),
]
for r_data in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(r_data):
        row[i].text = val
        if r_data[0].startswith("TOTAL"):
            for run in row[i].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph()
add_obs(doc, "A penalidade de Advertência corresponde a pontuações baixas; Suspensão 1–90 dias para pontuações mais elevadas.")
add_obs(doc, "O TAC é admissível quando a penalidade cabível é Advertência ou Suspensão de até determinado limite, conforme critérios de elegibilidade da Tela 8.")

# ── Salvar ────────────────────────────────────────────────────────────────────
output_path = "/home/user/desktop-tutorial/Arvore_Opcoes_Dosimetria_TAC.docx"
doc.save(output_path)
print(f"Documento salvo em: {output_path}")
