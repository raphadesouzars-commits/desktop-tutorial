#!/usr/bin/env python3
"""
Build Apostila LGPD/LAI v2 — parses the enriched markdown and outputs DOCX.
Reuses the full COGER visual system from the v1 script.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1E, 0x3A, 0x5F)
GOLD    = RGBColor(0xBF, 0xA0, 0x40)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
MUTED_C = RGBColor(0x70, 0x70, 0x70)
BLUE_D  = RGBColor(0x2E, 0x75, 0xB6)
GREEN   = RGBColor(0x1A, 0x6E, 0x1A)
GREEN_B = 'E8F5E9'

def rgb_hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

# ── XML helpers ──────────────────────────────────────────────────────────────
def set_shading_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.upper())
    for old in pPr.findall(qn('w:shd')): pPr.remove(old)
    pPr.append(shd)

def set_shading_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.upper())
    for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
    tcPr.append(shd)

def set_left_border(para, color, sz=18):
    pPr  = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None: pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), '6');    left.set(qn('w:color'), rgb_hex(color))
    for old in pBdr.findall(qn('w:left')): pBdr.remove(old)
    pBdr.append(left)

def set_cell_borders(cell, color='CCCCCC'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0');    el.set(qn('w:color'), color)
        tcBdr.append(el)
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    tcPr.append(tcBdr)

def add_runs(para, text, size=11, font='Calibri', color=None, base_bold=False, base_italic=False):
    parts = re.split(r'(\*\*|\*)', text)
    bold = base_bold; italic = base_italic
    for seg in parts:
        if seg == '**': bold   = not bold;   continue
        if seg == '*':  italic = not italic; continue
        if not seg: continue
        r = para.add_run(seg)
        r.font.name = font; r.font.size = Pt(size)
        r.font.bold = bold; r.font.italic = italic
        if color: r.font.color.rgb = color

# ── Style setup ──────────────────────────────────────────────────────────────
def setup_heading_styles(doc):
    def _style(name, font_name, size_pt, bold, color, space_before, space_after,
                left_border_color=None, left_border_sz=None):
        s = doc.styles[name]
        s.font.name      = font_name
        s.font.size      = Pt(size_pt)
        s.font.bold      = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(space_before)
        s.paragraph_format.space_after  = Pt(space_after)
        s.paragraph_format.keep_with_next = True
        pPr = s.element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None: pPr.remove(numPr)
        if left_border_color and left_border_sz:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None: pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), str(left_border_sz))
            left.set(qn('w:space'), '6')
            left.set(qn('w:color'), rgb_hex(left_border_color))
            for old in pBdr.findall(qn('w:left')): pBdr.remove(old)
            pBdr.append(left)

    _style('Heading 1', 'Arial Narrow', 20, True, NAVY, 4, 12)
    _style('Heading 2', 'Arial Narrow', 13, True, NAVY, 12, 5,
           left_border_color=GOLD, left_border_sz=18)
    _style('Heading 3', 'Calibri', 11, True, NAVY, 8, 3)

def insert_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)
    placeholder = OxmlElement('w:t')
    placeholder.text = 'Clique com o botão direito → Atualizar campo para ver o sumário com páginas.'
    run._r.append(placeholder)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)

# ── Components ────────────────────────────────────────────────────────────────
def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    set_left_border(p, GOLD, 12)
    r = p.add_run(text.upper())
    r.font.name = 'Arial Narrow'; r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = GOLD

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text.upper())
    r.font.name = 'Arial Narrow'; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY

def add_chapter(doc, number, title):
    p_eye = doc.add_paragraph()
    p_eye.paragraph_format.space_before = Pt(4); p_eye.paragraph_format.space_after = Pt(0)
    r_eye = p_eye.add_run(f'CAPÍTULO {number}')
    r_eye.font.name = 'Arial Narrow'; r_eye.font.size = Pt(9)
    r_eye.font.bold = True; r_eye.font.color.rgb = GOLD
    p2 = doc.add_paragraph(style='Heading 1')
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(title)
    r2.font.name = 'Arial Narrow'; r2.font.size = Pt(20)
    r2.font.bold = True; r2.font.color.rgb = NAVY
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0); sep.paragraph_format.space_after = Pt(10)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), rgb_hex(GOLD))
    pBdr.append(bot); pPr.append(pBdr)

def add_h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    set_left_border(p, GOLD, 18)
    r = p.add_run(text)
    r.font.name = 'Arial Narrow'; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = NAVY

def add_h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = NAVY

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)

def add_bullet(doc, text, indent=0.7):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(13.5); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p.add_run('— '); r0.font.name = 'Calibri'; r0.font.size = Pt(11)
    add_runs(p, text)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), rgb_hex(GOLD))
    pBdr.append(bot); pPr.append(pBdr)

def add_alert(doc, lines, kind='info', label=None):
    bg  = {'info': 'EBF3FB', 'warn': 'FFF8E1', 'danger': 'FDE8E8', 'green': GREEN_B, 'mono': 'F4F4F4'}
    bdc = {'info': BLUE_D, 'warn': GOLD, 'danger': RGBColor(0xC0, 0x00, 0x00), 'green': GREEN, 'mono': MUTED_C}
    bg_c = bg.get(kind, 'EBF3FB'); bd = bdc.get(kind, BLUE_D)
    if isinstance(lines, str): lines = [lines]
    first = True
    for line in lines:
        if not line.strip(): continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.4); p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(5) if first else Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_left_border(p, bd, 24); set_shading_para(p, bg_c)
        if first and label:
            r0 = p.add_run(f'{label}  ')
            r0.font.name = 'Arial Narrow'; r0.font.size = Pt(9); r0.font.bold = True; r0.font.color.rgb = bd
        add_runs(p, line, size=10)
        first = False
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def add_numbered(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.line_spacing = Pt(13.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r0 = p.add_run(f'{num}.  ')
    r0.font.name = 'Calibri'; r0.font.size = Pt(11); r0.font.bold = True; r0.font.color.rgb = NAVY
    add_runs(p, text)

def add_muted(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(9)
    r.font.italic = True; r.font.color.rgb = MUTED_C

def make_table(doc, headers, rows, col_widths, alt_bg='F5F8FF', header_bg=None):
    if header_bg is None: header_bg = rgb_hex(NAVY)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for j, (h, w) in enumerate(zip(headers, col_widths)):
        c = hdr.cells[j]; c.width = Cm(w)
        set_shading_cell(c, header_bg); set_cell_borders(c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        add_runs(p, h, size=9, color=WHITE, base_bold=True)
    for i, row in enumerate(rows):
        r = t.add_row()
        bg = alt_bg if i % 2 == 0 else 'FFFFFF'
        for j, (val, w) in enumerate(zip(row, col_widths)):
            c = r.cells[j]; c.width = Cm(w)
            set_shading_cell(c, bg); set_cell_borders(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            add_runs(p, val, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Markdown parser helpers ───────────────────────────────────────────────────
def get_alert_kind(label):
    if not label:
        return 'info'
    u = label.upper()
    if any(x in u for x in ['ATENÇÃO', 'PONTO-CHAVE', 'REGRA', 'SANÇÃO', 'HUMANIZAÇÃO',
                              'AVISO', 'ENUNCIADO', 'CUIDADO']):
        return 'warn'
    if any(x in u for x in ['ERRO', 'PERIGO', 'ESPECIAL']):
        return 'danger'
    if any(x in u for x in ['LEI ', 'ART. ', 'ART.', 'PORTARIA', 'CAPUT', 'LEI\n']):
        return 'mono'
    if label.startswith('LEI') or label.startswith('ART'):
        return 'mono'
    return 'info'

def parse_alert_block(raw_lines):
    """Extract label and content lines from a list of '> ...' lines."""
    content = []
    for line in raw_lines:
        if line.startswith('> '):
            content.append(line[2:])
        elif line.strip() == '>':
            content.append('')
        # else skip (shouldn't happen)

    label = None
    kind = 'info'
    if content:
        m = re.match(r'^\*\*([^\*]+)\*\*\s*(.*)', content[0])
        if m:
            label = m.group(1).strip()
            rest  = m.group(2).strip()
            content[0] = rest
            kind = get_alert_kind(label)

    # Remove leading/trailing empty strings
    while content and not content[0].strip():
        content.pop(0)
    while content and not content[-1].strip():
        content.pop()

    return label, kind, content

def parse_table_block(table_lines):
    """Parse markdown table lines into headers and rows."""
    if len(table_lines) < 2:
        return None, None

    def split_row(line):
        parts = line.strip().strip('|').split('|')
        return [p.strip() for p in parts]

    headers = split_row(table_lines[0])
    n_cols  = len(headers)
    rows = []
    for line in table_lines[2:]:  # skip separator line (|---|---|)
        if not line.strip() or line.strip().startswith('|---') or line.strip().startswith('|:--'):
            continue
        cells = split_row(line)
        # Pad or trim to match header count
        while len(cells) < n_cols: cells.append('')
        cells = cells[:n_cols]
        rows.append(cells)

    total_w = 16.0
    # Try to detect a narrow '#' or index column
    col_widths = []
    for h in headers:
        h_clean = h.strip('*').strip()
        if h_clean in ('#', 'N°', 'Nº', ''):
            col_widths.append(0.8)
        else:
            col_widths.append(None)

    fixed = sum(w for w in col_widths if w is not None)
    flex_count = sum(1 for w in col_widths if w is None)
    flex_w = (total_w - fixed) / flex_count if flex_count else total_w / n_cols

    col_widths = [w if w is not None else flex_w for w in col_widths]
    return headers, rows, col_widths

# ── Section-level header detection ───────────────────────────────────────────
RE_CHAPTER_EYE  = re.compile(r'^## CAPÍTULO (\d+)$')
RE_CHAPTER_TITLE = re.compile(r'^# (.+)$')
RE_H2_NUM       = re.compile(r'^## (\d+\.\S.*)')
RE_H3_NUM       = re.compile(r'^### (\d[\d\w\.-]*\s.+)')   # ### 1.7, ### 2.7, ### 3.2-A
RE_H3           = re.compile(r'^### (.+)')
RE_BULLET       = re.compile(r'^— (.+)')
RE_NUMBERED     = re.compile(r'^(\d+)\. (.+)')
RE_BOLD_NUMBERED = re.compile(r'^\*\*(\d+)\. (.+)')
RE_BOLD_HEADER  = re.compile(r'^\*\*([^*\n]+)\*\*$')

def is_separator(line): return line.strip() == '---'

def classify(line):
    """Return (type, data) for a markdown line."""
    s = line.rstrip()
    if not s.strip(): return ('blank', None)
    if is_separator(s): return ('sep', None)
    if s.startswith('> '): return ('alert_line', s)
    if s.strip() == '>': return ('alert_line', s)
    if s.startswith('|'): return ('table_line', s)
    m = RE_CHAPTER_EYE.match(s)
    if m: return ('chapter_eye', m.group(1))
    m = RE_CHAPTER_TITLE.match(s)
    if m: return ('chapter_title', m.group(1))
    if s.startswith('## EXERCÍCIOS PRÁTICOS'): return ('h1_exercises', None)
    if s.startswith('# EXERCÍCIOS PRÁTICOS'): return ('eyebrow_exercises', None)
    m = RE_H2_NUM.match(s)
    if m: return ('h2', m.group(1))
    if s.startswith('## '): return ('h2', s[3:].strip())
    m = RE_H3_NUM.match(s)
    if m: return ('h2', m.group(1))   # numbered ### → treat as h2
    m = RE_H3.match(s)
    if m: return ('h3', m.group(1))
    m = RE_BULLET.match(s)
    if m: return ('bullet', m.group(1))
    m = RE_BOLD_NUMBERED.match(s)
    if m: return ('bold_numbered', (m.group(1), m.group(2)))
    m = RE_NUMBERED.match(s)
    if m: return ('numbered', (m.group(1), m.group(2)))
    m = RE_BOLD_HEADER.match(s)
    if m: return ('bold_header', m.group(1))
    if s.startswith('*') and s.endswith('*') and not s.startswith('**'):
        return ('muted', s.strip('*'))
    return ('body', s)

# ── Main content parser ───────────────────────────────────────────────────────
def parse_content(doc, lines):
    i = 0
    pending_chapter_num = None

    while i < len(lines):
        line = lines[i]
        typ, data = classify(line)

        # ── Skip preamble / unused lines ─────────────────────────────────────
        if typ == 'blank':
            i += 1; continue

        # ── Separator ────────────────────────────────────────────────────────
        if typ == 'sep':
            add_separator(doc)
            i += 1; continue

        # ── Chapter eyebrow ──────────────────────────────────────────────────
        if typ == 'chapter_eye':
            pending_chapter_num = data
            i += 1; continue

        if typ == 'chapter_title':
            title = data
            page_break(doc)
            add_chapter(doc, pending_chapter_num or '?', title)
            pending_chapter_num = None
            i += 1; continue

        # ── Exercises header ─────────────────────────────────────────────────
        if typ == 'eyebrow_exercises':
            page_break(doc)
            add_eyebrow(doc, 'Exercícios Práticos com Gabarito Comentado')
            i += 1; continue

        if typ == 'h1_exercises':
            add_h1(doc, 'Exercícios Práticos')
            i += 1; continue

        # ── Headings ─────────────────────────────────────────────────────────
        if typ == 'h2':
            add_h2(doc, data)
            i += 1; continue

        if typ == 'h3':
            add_h3(doc, data)
            i += 1; continue

        # ── Alert block (accumulate all consecutive > lines) ─────────────────
        if typ == 'alert_line':
            raw = []
            while i < len(lines):
                t2, _ = classify(lines[i])
                if t2 == 'alert_line':
                    raw.append(lines[i].rstrip())
                    i += 1
                else:
                    break
            label, kind, content = parse_alert_block(raw)
            if content:
                add_alert(doc, content, kind=kind, label=label)
            continue

        # ── Table block ───────────────────────────────────────────────────────
        if typ == 'table_line':
            raw = []
            while i < len(lines):
                t2, _ = classify(lines[i])
                if t2 == 'table_line':
                    raw.append(lines[i].rstrip())
                    i += 1
                else:
                    break
            result = parse_table_block(raw)
            if result[0] is not None:
                headers, rows, col_widths = result
                make_table(doc, headers, rows, col_widths)
            continue

        # ── Bullet ───────────────────────────────────────────────────────────
        if typ == 'bullet':
            add_bullet(doc, data)
            i += 1; continue

        # ── Bold numbered (Solove-style: **1. Title** — text) ────────────────
        if typ == 'bold_numbered':
            num, rest = data
            add_numbered(doc, num, rest)
            i += 1; continue

        # ── Plain numbered list ───────────────────────────────────────────────
        if typ == 'numbered':
            num, rest = data
            add_numbered(doc, num, rest)
            i += 1; continue

        # ── Bold standalone header (e.g. **Gabarito comentado**) ─────────────
        if typ == 'bold_header':
            add_h3(doc, data)
            i += 1; continue

        # ── Muted (italic line) ───────────────────────────────────────────────
        if typ == 'muted':
            add_muted(doc, data)
            i += 1; continue

        # ── Body paragraph ────────────────────────────────────────────────────
        if typ == 'body':
            add_body(doc, data)
            i += 1; continue

        i += 1

# ════════════════════════════════════════════════════════════════════════════
# BUILD
# ════════════════════════════════════════════════════════════════════════════
def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin  = Cm(2.2);  sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5);  sec.right_margin  = Cm(2.5)
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.space_after = Pt(5)
    setup_heading_styles(doc)

    # ── CAPA ────────────────────────────────────────────────────────────────
    LOGO_PATH = '/tmp/logo_coger.png'
    try:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(20)
        p_logo.paragraph_format.space_after  = Pt(14)
        p_logo.add_run().add_picture(LOGO_PATH, width=Cm(7))
    except Exception as e:
        print(f'[WARN] Logo não inserida: {e}')

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after  = Pt(6)
    r = p_title.add_run('ACESSO À INFORMAÇÃO E PROTEÇÃO\nDE DADOS PESSOAIS NO ÂMBITO DISCIPLINAR')
    r.font.name = 'Arial Narrow'; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY

    add_separator(doc)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(4); p_sub.paragraph_format.space_after = Pt(4)
    r = p_sub.add_run('Apostila de treinamento — Versão enriquecida')
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = MUTED_C

    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub2.paragraph_format.space_after = Pt(2)
    r = p_sub2.add_run('Danielly Gontijo — Procuradora da Fazenda Nacional')
    r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.color.rgb = MUTED_C

    p_sub3 = doc.add_paragraph()
    p_sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub3.paragraph_format.space_after = Pt(16)
    r = p_sub3.add_run('Corregedoria da Receita Federal do Brasil — 2026')
    r.font.name = 'Calibri'; r.font.size = Pt(10); r.font.color.rgb = MUTED_C

    add_alert(doc, [
        '**Perguntas-chave desta oficina:**',
        '› Quando um dado disciplinar/correcional pode/deve ser publicizado — e quando isso caracteriza exposição irregular?',
        '› Como garantir que o tratamento de dados pessoais está sendo feito de forma legítima no âmbito disciplinar?',
    ], kind='info')

    page_break(doc)

    # ── SUMÁRIO ─────────────────────────────────────────────────────────────
    p_idx = doc.add_paragraph()
    p_idx.paragraph_format.space_before = Pt(0); p_idx.paragraph_format.space_after = Pt(14)
    r = p_idx.add_run('SUMÁRIO')
    r.font.name = 'Arial Narrow'; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = NAVY
    insert_toc(doc)
    add_separator(doc)

    # ── CONTEÚDO (parsed from markdown) ─────────────────────────────────────
    with open('/tmp/markdown_content.md', encoding='utf-8') as f:
        raw = f.read()

    # Strip the instruction prefix before the actual document
    raw = re.sub(r'^.*?## Apresentação', '## Apresentação', raw, count=1, flags=re.DOTALL)
    lines = raw.split('\n')

    page_break(doc)  # Page before Apresentação
    parse_content(doc, lines)

    return doc

doc = build()
OUT = '/tmp/Apostila_LGPD_LAI_ProcessosDisciplinares_COGER.docx'
doc.save(OUT)
print(f'SUCCESS: {OUT}')
print(f'Parágrafos: {len(doc.paragraphs)}  | Tabelas: {len(doc.tables)}')
