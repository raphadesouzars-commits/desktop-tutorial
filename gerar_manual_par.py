from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Cores institucionais ──────────────────────────────────────────────────────
AZUL_ESCURO = RGBColor(0x00, 0x20, 0x5B)   # --p281
VERDE       = RGBColor(0x64, 0xA7, 0x0B)   # --p369
CINZA_CLARO = RGBColor(0xF2, 0xF2, 0xF2)
BRANCO      = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_TEXTO = RGBColor(0x44, 0x44, 0x44)
AZUL_MEDIO  = RGBColor(0x1D, 0x3A, 0x78)
VERMELHO    = RGBColor(0xC0, 0x39, 0x2B)   # agravante
VERDE_ATEN  = RGBColor(0x27, 0xAE, 0x60)   # atenuante


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), kwargs.get(side, 'none'))
        if kwargs.get(side, 'none') != 'none':
            tag.set(qn('w:sz'), kwargs.get('sz', '4'))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_bottom_border_to_cell(cell, color='64A70B', sz='8'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '0')
    bot.set(qn('w:color'), color)
    tcBorders.append(bot)
    tcPr.append(tcBorders)


def no_space_before(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '60')
    pPr.append(spacing)


doc = Document()

# ── Margens ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ── Estilos base ─────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = CINZA_TEXTO


# ── Helpers ───────────────────────────────────────────────────────────────────
def heading1(text):
    """Título de seção numerada — fundo azul escuro."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, AZUL_ESCURO)
    for side in ('top','left','bottom','right'):
        try:
            set_cell_borders(cell, **{side: 'none'})
        except Exception:
            pass
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.font.name  = 'Arial'
    run.font.bold  = True
    run.font.size  = Pt(13)
    run.font.color.rgb = BRANCO
    doc.add_paragraph()  # espaço após
    return tbl


def heading2(text, cor=AZUL_MEDIO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Arial'
    run.font.bold  = True
    run.font.size  = Pt(11.5)
    run.font.color.rgb = cor
    # linha inferior verde
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '64A70B')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = 'Arial'
    run.font.bold  = True
    run.font.size  = Pt(11)
    run.font.color.rgb = AZUL_MEDIO
    return p


def body(text, bold_parts=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            before, _, remaining = remaining.partition(bp)
            if before:
                r = p.add_run(before)
                r.font.name = 'Arial'; r.font.size = Pt(11)
            r2 = p.add_run(bp)
            r2.font.name = 'Arial'; r2.font.size = Pt(11); r2.bold = True
        if remaining:
            r3 = p.add_run(remaining)
            r3.font.name = 'Arial'; r3.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name  = 'Arial'
        run.font.size  = Pt(11)
        run.italic     = italic
    return p


def cite(text):
    """Bloco de citação legal recuado."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x66)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p


def nota(text):
    """Nota de atenção com fundo cinza claro."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, CINZA_CLARO)
    add_bottom_border_to_cell(cell, color='64A70B', sz='6')
    p = cell.paragraphs[0]
    run = p.add_run('⚠  ' + text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x00)
    doc.add_paragraph()


def tabela(headers, rows, cor_header=AZUL_ESCURO, cor_texto_header=BRANCO, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # cabeçalho
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], cor_header)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.name = 'Arial'; run.font.bold = True
        run.font.size = Pt(10); run.font.color.rgb = cor_texto_header
    # linhas
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = CINZA_CLARO if ri % 2 == 0 else BRANCO
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = cells[ci].paragraphs[0].add_run(str(val))
            run.font.name = 'Arial'; run.font.size = Pt(10)
    # larguras
    if col_widths:
        for ri2, row2 in enumerate(tbl.rows):
            for ci2, cell2 in enumerate(row2.cells):
                if ci2 < len(col_widths):
                    cell2.width = Cm(col_widths[ci2])
    doc.add_paragraph()
    return tbl


def badge(text, cor=AZUL_MEDIO):
    """Rótulo colorido simulando o badge de etapa."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'  {text}  ')
    run.font.name  = 'Arial'
    run.font.bold  = True
    run.font.size  = Pt(10)
    run.font.color.rgb = BRANCO
    run.font.highlight_color = None
    # Usar shading via XML
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(cor[0], cor[1], cor[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    rPr.append(shd)
    return p


def page_break():
    doc.add_page_break()


def part_divider(title, subtitle=''):
    """Página divisória de parte do manual."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run(title)
    run.font.name = 'Arial'; run.font.bold = True; run.font.size = Pt(26)
    run.font.color.rgb = AZUL_ESCURO
    tbl_l = doc.add_table(rows=1, cols=1)
    tbl_l.alignment = WD_TABLE_ALIGNMENT.CENTER
    cl = tbl_l.rows[0].cells[0]
    set_cell_bg(cl, VERDE)
    cl.paragraphs[0].add_run(' ')
    cl.width = Cm(10)
    if subtitle:
        ps = doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ps.paragraph_format.space_before = Pt(20)
        rs = ps.add_run(subtitle)
        rs.font.name = 'Arial'; rs.font.size = Pt(13); rs.italic = True
        rs.font.color.rgb = AZUL_MEDIO
    page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  CAPA
# ══════════════════════════════════════════════════════════════════════════════
p_capa = doc.add_paragraph()
p_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_capa.paragraph_format.space_before = Pt(60)
r = p_capa.add_run('CORREGEDORIA DA RECEITA FEDERAL DO BRASIL')
r.font.name = 'Arial'; r.font.bold = True; r.font.size = Pt(13)
r.font.color.rgb = AZUL_ESCURO

p_capa2 = doc.add_paragraph()
p_capa2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_capa2.paragraph_format.space_before = Pt(40)
r2 = p_capa2.add_run('MANUAL TEÓRICO-PRÁTICO')
r2.font.name = 'Arial'; r2.font.bold = True; r2.font.size = Pt(22)
r2.font.color.rgb = AZUL_ESCURO

p_capa3 = doc.add_paragraph()
p_capa3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_capa3.add_run('Calculadora de Multa PAR')
r3.font.name = 'Arial'; r3.font.bold = True; r3.font.size = Pt(18)
r3.font.color.rgb = VERDE

# linha decorativa
tbl_linha = doc.add_table(rows=1, cols=1)
tbl_linha.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_linha = tbl_linha.rows[0].cells[0]
set_cell_bg(cell_linha, VERDE)
cell_linha.paragraphs[0].add_run(' ')
cell_linha.width = Cm(12)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(30)
r_sub = p_sub.add_run('Processo Administrativo de Responsabilização — PAR\nLei nº 12.846/2013 | Decreto nº 11.129/2022')
r_sub.font.name = 'Arial'; r_sub.font.size = Pt(12)
r_sub.font.color.rgb = AZUL_MEDIO

p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub2.paragraph_format.space_before = Pt(60)
r_sub2 = p_sub2.add_run('Versão 2026 — Conforme Manual de Responsabilização de Entes Privados da CGU (3ª ed., jun/2026)')
r_sub2.font.name = 'Arial'; r_sub2.font.size = Pt(11)
r_sub2.font.color.rgb = CINZA_TEXTO

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  APRESENTAÇÃO
# ══════════════════════════════════════════════════════════════════════════════
heading1('APRESENTAÇÃO')
body('Este manual tem dupla finalidade. Na Parte I, apresenta os fundamentos teóricos do Processo Administrativo de Responsabilização (PAR): o fenômeno da corrupção e suas distorções, a origem e a lógica da Lei Anticorrupção, os sujeitos alcançados, os atos lesivos, o rito processual, o julgamento, os recursos, o acordo de leniência e os cadastros de sanções. Na Parte II, ensina a dosimetria da multa e o uso prático da Calculadora de Multa PAR, etapa por etapa, com fórmulas, tabelas legais, casos especiais e estudos de caso comentados.')
body('O operador que domina apenas a mecânica da calculadora, sem compreender a teoria que a sustenta, corre o risco de aplicar percentuais sem o necessário juízo crítico. Por isso, recomenda-se a leitura integral da Parte I antes de manusear a ferramenta descrita na Parte II.')

# ══════════════════════════════════════════════════════════════════════════════
#  PARTE I — FUNDAMENTOS TEÓRICOS
# ══════════════════════════════════════════════════════════════════════════════
part_divider('PARTE I', 'Fundamentos Teóricos do PAR')

# ── I.1 A CORRUPÇÃO E SUAS DISTORÇÕES ─────────────────────────────────────────
heading1('1. A CORRUPÇÃO E SUAS DISTORÇÕES')
body('Compreender a finalidade da Lei Anticorrupção exige, antes, compreender o mal que ela combate. A metáfora tradicional de que "a corrupção seria o lubrificante da burocracia" — sugerindo que a propina apenas agilizaria processos emperrados — é falsa e perigosa. A corrupção não lubrifica; ela corrói.')
heading2('1.1 Tipos de Corrupção')
tabela(['Tipo', 'Descrição', 'Consequência'],
    [
        ['"Speedy money" (dinheiro de velocidade)', 'Propina paga para acelerar processos burocráticos; o agente público usa seu poder para remover entraves que ele mesmo cria', 'Gera tolerância generalizada, fomenta novas propinas e acentua as distorções econômicas'],
        ['"Rent seeking" (busca de renda)', 'Atuação em licenças ou mercados escassos; a ideia de que "só o eficiente paga propina" é falsa', 'Confere vantagem competitiva exclusiva a quem detém a melhor relação espúria com agentes corruptos'],
    ],
    col_widths=[3.5, 8, 5])
heading2('1.2 Distorções Econômicas e Sociais')
bullet('Seleção adversa: nem sempre vencem as empresas mais eficientes, mas as dispostas a corromper;')
bullet('Alocação ineficiente de recursos: desperdício em obras e serviços superfaturados, desnecessários e de qualidade inferior;')
bullet('Redução da competição: empresas honestas afastam-se dos "pseudoleilões" públicos, gerando má alocação de talentos.')
nota('A dosimetria não é mero exercício aritmético. Cada agravante e atenuante traduz, em percentual, o grau de reprovabilidade da conduta à luz desses danos sociais.')

# ── I.2 A LEI ANTICORRUPÇÃO ────────────────────────────────────────────────────
heading1('2. A LEI ANTICORRUPÇÃO: ORIGEM E CARACTERÍSTICAS')
heading2('2.1 Contexto Histórico e Pressão Internacional')
body('A Lei nº 12.846/2013 nasceu de compromissos internacionais: a Convenção da OCDE sobre o Combate da Corrupção de Funcionários Públicos Estrangeiros, a Convenção da ONU contra a Corrupção (Convenção de Mérida) e a Convenção da OEA.')
body('Historicamente, o foco criminal punia apenas a pessoa física do corruptor e do corrompido. A Lei Anticorrupção muda o paradigma para responsabilizar diretamente a pessoa jurídica — ou seja, atinge o lado da oferta da corrupção. As manifestações de junho de 2013 aceleraram sua sanção.')
heading2('2.2 Características Centrais')
tabela(['Característica', 'Conteúdo'],
    [
        ['Responsabilidade objetiva da PJ', 'Dispensa dolo ou culpa. Basta o nexo causal entre o ato lesivo e o benefício (direto ou indireto) auferido pela empresa'],
        ['Esfera administrativa e cível', 'A responsabilização ocorre de forma independente e simultânea em ambas as esferas'],
        ['Independência das instâncias', 'A punição da PJ não exclui a responsabilidade individual de dirigentes, administradores ou de qualquer pessoa física'],
    ],
    col_widths=[4.5, 12])
nota('Como a responsabilidade é objetiva, o PAR não discute a "intenção" da empresa. A Parte II pressupõe que o mérito (ato lesivo + nexo causal) já foi reconhecido; a dosimetria apenas quantifica a sanção.')

# ── I.3 SUJEITOS PASSIVOS ──────────────────────────────────────────────────────
heading1('3. SUJEITOS PASSIVOS E ALCANCE DA LEI')
heading2('3.1 A Quem se Aplica (Art. 1º, parágrafo único, da LAC)')
body('A lei alcança as seguintes entidades, personificadas ou não, de direito ou de fato:')
bullet('Sociedades empresárias e sociedades simples;')
bullet('Fundações e associações;')
bullet('Sociedades estrangeiras com sede, filial ou representação no território brasileiro.')
heading2('3.2 Sucessão Empresarial (Art. 4º)')
body('A responsabilidade subsiste em caso de alteração contratual, transformação, incorporação, fusão ou cisão.')
tabela(['Evento', 'Efeito na responsabilidade'],
    [
        ['Alteração contratual / transformação', 'A PJ é a mesma — responsabilidade inalterada'],
        ['Incorporação / fusão', 'Responsabilidade da sucessora restrita ao limite do patrimônio transferido, salvo simulação ou fraude evidente'],
    ],
    col_widths=[5, 12])
heading2('3.3 Desconsideração da Personalidade Jurídica (Art. 14)')
body('A personalidade jurídica pode ser desconsiderada quando usada com abuso de direito para facilitar, encobrir ou dissimular o ato lesivo, ou quando houver confusão patrimonial. Os efeitos das sanções estendem-se aos sócios e administradores com poderes de gerência.')

# ── I.4 ATOS LESIVOS ───────────────────────────────────────────────────────────
heading1('4. OS ATOS LESIVOS À ADMINISTRAÇÃO PÚBLICA (Art. 5º)')
body('A correta identificação da espécie e do número de condutas é essencial para o agravante de concurso de atos lesivos (Art. 22, I).')
heading2('4.1 Vantagens Indevidas e Fraudes Gerais')
bullet('Art. 5º, I — Prometer, oferecer ou dar vantagem indevida a agente público ou a terceiro a ele relacionado;')
bullet('Art. 5º, II — Financiar, custear, patrocinar ou subvencionar a prática dos atos ilícitos;')
bullet('Art. 5º, III — Utilizar interposta pessoa para ocultar interesses ou identidade dos beneficiários.')
heading2('4.2 Licitações e Contratos (Art. 5º, IV)')
bullet('Fraudar ou frustrar o caráter competitivo de licitação;')
bullet('Impedir, perturbar ou fraudar atos do procedimento licitatório;')
bullet('Afastar licitante por fraude ou vantagem; fraudar licitação ou contrato dela decorrente;')
bullet('Criar PJ de modo fraudulento para participar de licitação ou contrato;')
bullet('Obter vantagem indevida de modificações/prorrogações; manipular o equilíbrio econômico-financeiro.')
heading2('4.3 Obstrução à Fiscalização (Art. 5º, V)')
body('Dificultar investigação ou fiscalização de órgãos, entidades ou agentes públicos, ou intervir em sua atuação — inclusive agências reguladoras e órgãos de fiscalização do sistema financeiro.')
nota('No Art. 22, I, as colunas representam o número de espécies distintas de atos do Art. 5º; as linhas, o número de condutas. Oferta de vantagem (inc. I) + obstrução (inc. V) = duas espécies distintas.')

# ── I.5 O RITO DO PAR ──────────────────────────────────────────────────────────
heading1('5. O PROCESSO ADMINISTRATIVO DE RESPONSABILIZAÇÃO — RITO')
body('O PAR divide-se em duas macrofases: a Fase Interna (admissibilidade e investigação) e a Fase Externa (contraditório e julgamento).')
heading2('5.1 Macrofase I — Admissibilidade e Investigação Preliminar')
bullet('Notícia do fato: denúncias, auditorias, representações fiscais ou operações policiais;')
bullet('Investigação Preliminar (IP): inquisitorial e sigilosa; busca indícios de autoria e materialidade;')
bullet('Juízo de admissibilidade: arquivamento fundamentado, diligências/IP complementar, ou instauração do PAR.')
heading2('5.2 Macrofase II — Fase Externa e Comissão')
bullet('Instauração por Portaria publicada no DOU; comissão de 2 ou mais servidores estáveis;')
bullet('Prazo padrão de 180 dias, prorrogável mais de uma vez (Enunciado CGU nº 24/2019);')
bullet('Instrução: provas documentais, testemunhais, perícias e requisição de informações.')
heading2('5.3 Indiciação, Citação e Defesa')
tabela(['Ato', 'Descrição'],
    [
        ['Nota de Indiciação', 'Peça técnica que individualiza e descreve as condutas, capitulando os atos lesivos com base nas provas'],
        ['Citação', 'Intimação formal para defesa; prazo de Defesa Escrita de 30 dias da citação'],
        ['Revelia', 'Se a PJ citada não se defende, prossegue o processo, assegurado o direito de intervir depois'],
    ],
    col_widths=[4, 13])
nota('A tempestividade da admissão de responsabilidade (Art. 23, IV) é medida em relação a esses marcos: antes da instauração, no prazo de defesa (30 dias), nas alegações finais, ou após.')

# ── I.6 JULGAMENTO ─────────────────────────────────────────────────────────────
heading1('6. JULGAMENTO, SANÇÕES E FASE RECURSAL')
heading2('6.1 Relatório Final e Julgamento')
body('Concluída a instrução, a comissão elabora Relatório Final, opinando sobre a ocorrência do ato lesivo, a responsabilidade da PJ e a proposta de dosimetria — é aqui que a Calculadora é instrumental. O processo segue para parecer jurídico obrigatório (controle de legalidade) e julgamento.')
nota('No âmbito da Receita Federal, a competência para julgar o PAR é delegada ao Corregedor da RFB (no âmbito federal geral, a competência máxima é do Ministro de Estado).')
heading2('6.2 Sanções Aplicáveis (cumulativas)')
bullet('Multa pecuniária — conforme a dosimetria (Parte II);')
bullet('Publicação extraordinária da decisão condenatória — a expensas da PJ: DOU, sítio do órgão lesado, mídia de grande circulação (mín. 1 dia) e painel no estabelecimento (30 a 120 dias).')
heading2('6.3 Fase Recursal')
tabela(['Item', 'Regra'],
    [
        ['Pedido de Reconsideração', 'Prazo de 10 dias da publicação no DOU'],
        ['Efeito', 'Suspensivo — suspende os efeitos da decisão'],
        ['Cumprimento', 'Confirmada a condenação, 30 dias para pagar a multa e comprovar a publicação'],
    ],
    col_widths=[5, 12])

# ── I.7 LENIÊNCIA E CADASTROS ──────────────────────────────────────────────────
heading1('7. ACORDO DE LENIÊNCIA E CADASTROS DE SANÇÕES')
heading2('7.1 Acordo de Leniência (Art. 16)')
body('A autoridade máxima do órgão pode celebrar acordo com a PJ que colabore efetivamente. Requisitos cumulativos:')
bullet('Ser a primeira a manifestar interesse em cooperar (quando aplicável o critério de precedência);')
bullet('Cessação completa do envolvimento na infração a partir da propositura;')
bullet('Admissão integral da responsabilidade objetiva;')
bullet('Cooperação plena, contínua e de boa-fé.')
body('Benefícios:')
bullet('Redução de até 2/3 do valor da multa aplicável;')
bullet('Isenção da publicação extraordinária;')
bullet('Isenção da proibição de receber incentivos, subsídios, subvenções, doações ou empréstimos públicos.')
nota('Na leniência, o prazo de 5 anos da reincidência (Art. 22, V) conta da celebração do acordo até cinco anos após a declaração de seu cumprimento.')
heading2('7.2 Cadastros Nacionais de Sanções')
tabela(['Cadastro', 'Finalidade'],
    [
        ['CGU-PJ / Siasg', 'Controle interno dos processos de responsabilização de entes privados'],
        ['CEIS — Empresas Inidôneas e Suspensas', 'Consolida sanções de restrição ao direito de licitar e contratar de todos os entes federativos'],
        ['CNEP — Empresas Punidas', 'Dá publicidade às sanções aplicadas com base na Lei nº 12.846/2013'],
    ],
    col_widths=[5.5, 11])

# ══════════════════════════════════════════════════════════════════════════════
#  PARTE II — DOSIMETRIA E USO DA CALCULADORA
# ══════════════════════════════════════════════════════════════════════════════
part_divider('PARTE II', 'Dosimetria e Uso da Calculadora')

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUÇÃO E FINALIDADE
# ══════════════════════════════════════════════════════════════════════════════
heading1('1. INTRODUÇÃO E FINALIDADE')
body('A Calculadora de Multa PAR é uma ferramenta de apoio à decisão desenvolvida pela Corregedoria da Receita Federal do Brasil para uso nos Processos Administrativos de Responsabilização (PAR) instaurados com base na Lei Anticorrupção (Lei nº 12.846/2013) e seu decreto regulamentador (Decreto nº 11.129/2022).')
body('A ferramenta tem como objetivos:')
bullet('Padronizar a dosimetria das sanções pecuniárias aplicadas às pessoas jurídicas;')
bullet('Transparência — cada percentual aplicado está vinculado a um dispositivo legal específico;')
bullet('Rastreabilidade — o relatório final gerado documenta todos os parâmetros utilizados, servindo como instrumento de instrução processual;')
bullet('Reduzir erros de cálculo, especialmente nas interações entre os limites mínimo, máximo, vantagem auferida e alíquota preliminar.')
body('A calculadora funciona integralmente offline (sem dependência de servidores externos), sendo adequada para uso em ambientes com restrições de conectividade.')

# ══════════════════════════════════════════════════════════════════════════════
#  2. FUNDAMENTO LEGAL
# ══════════════════════════════════════════════════════════════════════════════
heading1('2. FUNDAMENTO LEGAL')
heading2('2.1 Lei nº 12.846/2013 — Lei Anticorrupção (LAC)')
tabela(
    ['Dispositivo', 'Conteúdo'],
    [
        ['Art. 5º',      'Tipificação dos atos lesivos à administração pública'],
        ['Art. 6º, I',   'Multa de 0,1% a 20% do faturamento bruto da pessoa jurídica'],
        ['Art. 6º, II',  'Publicação extraordinária da decisão condenatória'],
    ],
    col_widths=[4, 12]
)
heading2('2.2 Decreto nº 11.129/2022')
tabela(
    ['Dispositivo', 'Conteúdo'],
    [
        ['Art. 20',  'Base de cálculo: faturamento bruto do último exercício anterior ao PAR'],
        ['Art. 21',  'Hipótese subsidiária: ausência de faturamento no exercício anterior'],
        ['Art. 22',  'Fatores agravantes (I a VI)'],
        ['Art. 23',  'Fatores atenuantes (I a V)'],
        ['Art. 24',  'Publicação extraordinária e critérios de prazo'],
        ['Art. 25',  'Limites mínimo e máximo da multa'],
        ['Art. 26',  'Métodos de apuração da vantagem auferida (I a III)'],
    ],
    col_widths=[4, 12]
)
heading2('2.3 Orientações complementares')
bullet('IN CGU nº 1/2015 — metodologia de avaliação do programa de integridade;')
bullet('Manual de Responsabilização de Entes Privados da CGU (3ª ed., jun/2026) — metodologia oficial de dosimetria (Tabela 3.1: interrupção de serviço público; Tabela 6: colaboração);')
bullet('Portaria Conjunta CGU nº 6/2022 — Adendo nº 1 ao Manual Prático de Avaliação de Programa de Integridade; fator multiplicador de 1,25 (Art. 23, V);')
bullet('Guia CGU de Responsabilização de Pessoas Jurídicas — parâmetros interpretativos e tabelas de referência;')
bullet('Tabela 5 do Guia CGU — escalonamento dos percentuais de ressarcimento (Art. 23, II).')

# ══════════════════════════════════════════════════════════════════════════════
#  3. ESTRUTURA GERAL
# ══════════════════════════════════════════════════════════════════════════════
heading1('3. ESTRUTURA GERAL DA CALCULADORA')
body('A calculadora é organizada em 14 etapas sequenciais (Steps 0 a 13), agrupadas em três blocos funcionais:')
tabela(
    ['Bloco', 'Etapas', 'Conteúdo'],
    [
        ['A — Base de Cálculo',   'Etapas 0 e 1',  'Faturamento Bruto e Vantagens Auferida/Pretendida'],
        ['B — Dosimetria',        'Etapas 2 a 12', 'Agravantes (Art. 22) e Atenuantes (Art. 23)'],
        ['C — Resultado',         'Etapa 13',      'Publicação Extraordinária e Resultado Final (Art. 25)'],
    ],
    col_widths=[5, 4, 8]
)

# ══════════════════════════════════════════════════════════════════════════════
#  4. ETAPA 0 — BASE DE CÁLCULO
# ══════════════════════════════════════════════════════════════════════════════
heading1('4. ETAPA 0 — BASE DE CÁLCULO: FATURAMENTO BRUTO')
heading2('4.1 Fundamento Legal')
cite('Art. 20 do Decreto nº 11.129/2022. A multa prevista no inciso I do caput do art. 6º da Lei nº 12.846/2013 terá como base de cálculo o faturamento bruto da pessoa jurídica no último exercício anterior ao da instauração do PAR, excluídos os tributos.')
cite('Art. 21. Caso a pessoa jurídica comprovadamente não tenha tido faturamento no último exercício anterior ao da instauração do PAR, deve-se considerar como base de cálculo da multa o valor do último faturamento bruto apurado pela pessoa jurídica, excluídos os tributos incidentes sobre vendas, que terá seu valor atualizado pelo IPCA até o último dia do exercício anterior ao da instauração do PAR.')
heading2('4.2 Fluxo de Decisão')
tabela(
    ['Situação', 'Regra Aplicável', 'Fundamento'],
    [
        ['PJ teve faturamento no exercício anterior ao PAR', 'Usar faturamento do exercício de referência (Ano PAR − 1)', 'Art. 20'],
        ['PJ não teve faturamento no exercício anterior',    'Usar o último faturamento bruto apurado, atualizado pelo IPCA', 'Art. 21'],
    ],
    col_widths=[6, 7, 4]
)
heading2('4.3 Modos de Apuração')
tabela(
    ['Modo', 'Descrição', 'Fundamento'],
    [
        ['DRE (Regime Geral)',         'Receita Bruta − Tributos sobre Vendas', 'Art. 20, caput'],
        ['Simples Nacional (ME/EPP)',  'Receita bruta total apurada nas declarações do Simples', 'LC nº 123/2006, art. 3º, §1º'],
        ['Estimativa pela autoridade','Hipótese residual — outros elementos quando não há dados contábeis', 'Art. 20, §1º, III'],
    ],
    col_widths=[5, 8, 4]
)
heading2('4.4 Deduções: Atenção à DRE')
body('São dedutíveis apenas os tributos incidentes sobre as vendas: PIS, COFINS, IPI, ICMS e ISS.')
nota('Ponto crítico: na DRE contábil, a receita bruta é deduzida também de abatimentos, devoluções e vendas canceladas. Para fins de PAR, contudo, APENAS os tributos têm previsão legal de exclusão. Devoluções e vendas canceladas NÃO reduzem a base de cálculo da multa. Atenção redobrada ao importar números diretamente da DRE.')
heading2('4.5 Critérios de Exceção (faturamento não apurável)')
tabela(['Ordem', 'Critério', 'Fundamento'],
    [
        ['1º', 'Faturamento bruto do último ano com receita, atualizado pelo IPCA', 'Art. 21'],
        ['2º', 'Faturamento estimado: mercado, capital social, dados setoriais fundamentados', 'Art. 20, §1º, III'],
        ['3º', 'Valor absoluto fixo legal: multa de R$ 6 mil a R$ 60 milhões', 'Art. 6º, I, da Lei'],
    ],
    col_widths=[2, 11, 4])
nota('Como a calculadora funciona offline, a atualização pelo IPCA (Art. 21) deve ser realizada previamente na Calculadora do Cidadão (BCB/IPCA) ou na Calculadora de Multa da CGU. Informe o valor já corrigido no campo correspondente.')

# ══════════════════════════════════════════════════════════════════════════════
#  5. ETAPA 1 — VANTAGENS AUFERIDA E PRETENDIDA
# ══════════════════════════════════════════════════════════════════════════════
heading1('5. ETAPA 1 — VANTAGENS AUFERIDA E PRETENDIDA (Art. 26)')
heading2('5.1 Distinção Legal')
tabela(
    ['Conceito', 'Definição', 'Função na Multa', 'Fundamento'],
    [
        ['Vantagem Auferida',   'Ganho real e efetivamente obtido com o ato lesivo',     'Piso mínimo inafastável da multa',                     'Art. 25, I'],
        ['Vantagem Pretendida', 'Ganho planejado/visado, ainda que não concretizado',    'Referência para o teto máximo (maior entre as duas)',   'Art. 25, II'],
    ],
    col_widths=[4.5, 6, 5.5, 3]
)
heading2('5.2 Métodos de Apuração da Vantagem Auferida')
heading3('Modo Direto')
body('Informar o valor já apurado externamente. Utilizar zero quando a vantagem não é estimável.')
heading3('Art. 26, I — Diferença entre receita auferida e custos lícitos')
body('Aplicável quando há receita auferida indevidamente (ex.: contratos obtidos por ato ilícito).')
cite('Vantagem = Receita Auferida − CMV/CSP − Despesas Lícitas − IRPJ/CSLL')
body('Suporta múltiplos exercícios com tratamento de prejuízos acumulados.')
heading3('Art. 26, II — Custos e despesas não suportados')
body('Valor correspondente aos custos e despesas que a PJ deixou de suportar. Exemplos: tributos não recolhidos, multas regulatórias evitadas, encargos não pagos.')
heading3('Art. 26, III — Acréscimo patrimonial por ação/omissão do Poder Público')
body('Acréscimo patrimonial que não teria ocorrido sem o ato lesivo. Exemplos: juros subsidiados obtidos indevidamente, uso de informação privilegiada.')
nota('Não confundir vantagem auferida com pretendida. Inflar o piso mínimo com a vantagem pretendida é erro metodológico. A calculadora trata cada valor na função legal correta.')

# ══════════════════════════════════════════════════════════════════════════════
#  6. AGRAVANTES
# ══════════════════════════════════════════════════════════════════════════════
heading1('6. AGRAVANTES — ART. 22, I A VI')
heading2('6.1 Art. 22, I — Concurso de Atos Lesivos', cor=VERMELHO)
body('Cruzamento entre quantidade de condutas ilícitas (linhas) e número de espécies distintas de atos lesivos do art. 5º da LAC (colunas). Aplica-se o percentual da célula correspondente ao caso concreto.')
tabela(
    ['Condutas ↓ / Espécies →', '1 espécie', '2 espécies', '3 espécies', '4 ou mais'],
    [
        ['1 conduta (ato isolado)', '—',    '0,5%', '1,0%', '1,5%'],
        ['2 condutas',             '0,5%', '1,0%', '1,5%', '2,0%'],
        ['3 condutas',             '1,0%', '1,5%', '2,0%', '2,5%'],
        ['4 condutas',             '1,5%', '2,0%', '2,5%', '3,0%'],
        ['5 condutas',             '2,0%', '2,5%', '3,0%', '3,5%'],
        ['6 condutas',             '2,5%', '3,0%', '3,5%', '4,0%'],
        ['7 ou mais',              '3,0%', '3,5%', '4,0%', '4,0%'],
    ],
    cor_header=VERMELHO,
    col_widths=[5.5, 3, 3, 3, 3]
)

heading2('6.2 Art. 22, II — Ciência/Tolerância Hierárquica', cor=VERMELHO)
body('O percentual varia conforme a posição hierárquica do agente com conhecimento ou tolerância, medida a partir do topo da estrutura.')
tabela(
    ['Nível hierárquico', '%'],
    [
        ['Ausência de conhecimento pelo corpo diretivo e gerencial',          '0%'],
        ['Tolerância/ciência do 5º nível abaixo dos administradores',         '1,0%'],
        ['Tolerância/ciência do 4º nível abaixo dos administradores',         '1,5%'],
        ['Tolerância/ciência do 3º nível abaixo dos administradores',         '2,0%'],
        ['Tolerância/ciência do 2º nível (imediatamente abaixo)',             '2,5%'],
        ['Tolerância/ciência dos sócios, acionistas ou administradores',      '3,0%'],
    ],
    cor_header=VERMELHO,
    col_widths=[13, 3]
)

heading2('6.3 Art. 22, III — Interrupção/Descumprimento (três hipóteses)', cor=VERMELHO)
body('Avalie cada uma das três hipóteses abaixo. Aplica-se automaticamente o maior percentual entre elas.')
heading3('a) Interrupção no Fornecimento de Serviço Público')
tabela(
    ['Situação', '%'],
    [
        ['Ausência de interrupção no fornecimento de serviço público',                                   '0%'],
        ['Até 1 semana, ou impacto em município com até 100 mil habitantes',                             '1,0%'],
        ['Até 2 semanas, ou impacto em município com até 400 mil habitantes',                            '2,0%'],
        ['Até 3 semanas, ou município com mais de 400 mil hab., ou mais de um município do mesmo Estado', '3,0%'],
        ['Superior a 4 semanas, ou dois ou mais Estados, ou dois ou mais municípios com +400 mil hab.',  '4,0%'],
    ],
    cor_header=VERMELHO,
    col_widths=[13, 3]
)
nota('Atualização do Manual CGU (3ª ed., jun/2026), Tabela 3.1: os limiares de porte de município passam a ser 100 mil e 400 mil habitantes, substituindo os antigos critérios "vila/povoado" e de 500 mil habitantes.')
heading3('b) Interrupção na Execução de Obra Contratada')
tabela(
    ['Período ↓ / Residual →', '< 10%', 'Até 30%', 'Até 50%', 'Até 70%', '> 70%'],
    [
        ['Até 6 meses', '0,5%', '1,0%', '1,5%', '2,0%', '2,5%'],
        ['Até 1 ano',   '1,0%', '1,5%', '2,0%', '2,5%', '3,0%'],
        ['Até 2 anos',  '1,5%', '2,0%', '2,5%', '3,0%', '3,5%'],
        ['+ 2 anos',    '2,0%', '2,5%', '3,0%', '3,5%', '4,0%'],
    ],
    cor_header=VERMELHO,
    col_widths=[4.5, 2.5, 2.5, 2.5, 2.5, 2.5]
)
heading3('c) Descumprimento de Requisitos Regulatórios')
tabela(
    ['Situação', '%'],
    [
        ['Ausência de descumprimento',          '0%'],
        ['Parcial, com prestação do serviço',   '1,0%'],
        ['Total, com prestação do serviço',     '2,0%'],
        ['Parcial, sem prestação do serviço',   '3,0%'],
        ['Total, sem prestação do serviço',     '4,0%'],
    ],
    cor_header=VERMELHO,
    col_widths=[13, 3]
)

heading2('6.4 Art. 22, IV — Situação Econômica do Infrator', cor=VERMELHO)
body('Aplica-se 1% fixo quando a PJ apresenta, cumulativamente, os três indicadores:')
tabela(
    ['Indicador', 'Fórmula'],
    [
        ['Solvência geral > 1',  '(Ativo Circulante + Ativo Não Circulante) ÷ (Passivo Circulante + Passivo Não Circulante)'],
        ['Liquidez geral > 1',   '(Ativo Circulante + Ativo Realizável a LP) ÷ (Passivo Circulante + Passivo Não Circulante)'],
        ['Lucro líquido positivo', 'Último exercício anterior ao PAR'],
    ],
    col_widths=[5, 12]
)
tabela(['Situação', '%'],
    [['Algum indicador não atendido', '0%'], ['Todos os três atendidos cumulativamente', '1,0%']],
    cor_header=VERMELHO, col_widths=[13, 3])

heading2('6.5 Art. 22, V — Reincidência', cor=VERMELHO)
body('Ocorrência de nova infração tipificada como ato lesivo (art. 5º da LAC), idêntica ou não, em menos de cinco anos contados da publicação do julgamento anterior. No caso de acordo de leniência, o prazo conta da data da celebração até cinco anos após a declaração de cumprimento.')
tabela(['Situação', '%'],
    [['Não reincidente', '0%'], ['Sim — nova infração em menos de 5 anos', '3,0%']],
    cor_header=VERMELHO, col_widths=[13, 3])

heading2('6.6 Art. 22, VI — Valor dos Contratos com o Ente Lesado', cor=VERMELHO)
body('Somatório dos contratos, convênios, acordos e demais instrumentos mantidos ou pretendidos com o órgão/entidade lesado nos anos da prática do ato lesivo.')
tabela(['Somatório dos instrumentos', '%'],
    [
        ['Até R$ 500 mil',                          '0%'],
        ['Superior a R$ 500 mil até R$ 1,5 milhão', '1,0%'],
        ['Superior a R$ 1,5 milhão até R$ 10 milhões', '2,0%'],
        ['Superior a R$ 10 milhões até R$ 50 milhões', '3,0%'],
        ['Superior a R$ 50 milhões até R$ 250 milhões','4,0%'],
        ['Superior a R$ 250 milhões',               '5,0%'],
    ],
    cor_header=VERMELHO, col_widths=[13, 3])

# ══════════════════════════════════════════════════════════════════════════════
#  7. ATENUANTES
# ══════════════════════════════════════════════════════════════════════════════
heading1('7. ATENUANTES — ART. 23, I A V')
heading2('7.1 Art. 23, I — Não Consumação', cor=VERDE_ATEN)
body('Aplica-se quando o ato lesivo não chegou a produzir seus efeitos plenos (tentativa).')
tabela(['Situação', '%'],
    [['Ato lesivo consumado', '0%'], ['Ato lesivo não consumado (tentativa)', '−0,5%']],
    cor_header=VERDE_ATEN, col_widths=[13, 3])

heading2('7.2 Art. 23, II — Ressarcimento e Devolução Espontânea', cor=VERDE_ATEN)
body('Atenuante de até 1,0% pela devolução espontânea da vantagem auferida e/ou ressarcimento dos danos, conforme Tabela 5 do Guia CGU.')
tabela(['Situação', '%'],
    [
        ['Ausência de devolução e de ressarcimento', '0%'],
        ['Devolução da vantagem sem ressarcimento dos danos; ou ressarcimento sem devolução da vantagem', '−0,5%'],
        ['Devolução da vantagem e ressarcimento dos danos', '−1,0%'],
        ['Devolução e inexistência/falta de comprovação de danos; ou ressarcimento e ausência de estimativa da vantagem', '−1,0%'],
        ['Inexistência ou falta de comprovação de vantagem e de danos', '−1,0%'],
    ],
    cor_header=VERDE_ATEN, col_widths=[13, 3])
nota('O percentual máximo de 1,0% exige confirmação expressa de devolução integral. Devolução parcial reduz o percentual para 0,5%.')

heading2('7.3 Art. 23, III — Colaboração', cor=VERDE_ATEN)
body('Reconhecida ainda que não haja admissão de responsabilidade. Conforme a Tabela 6 do Manual CGU (3ª ed., jun/2026), a colaboração NÃO é soma fixa de 0,5% por condição, mas valoração por faixa:')
tabela(['Condições factualmente presentes', 'Percentual'],
    [
        ['Nenhuma condição', '0%'],
        ['Uma ou duas condições', 'Faixa discricionária de 0,5% a 1,0% (fixada pela autoridade conforme a utilidade e relevância da colaboração)'],
        ['As três condições, simultaneamente', '1,5% fixo (sem discricionariedade)'],
    ],
    cor_header=VERDE_ATEN, col_widths=[5, 11])
body('As três condições que se registram (fato objetivo): (i) admitiu a prática do ato; (ii) forneceu elementos para a apuração; (iii) renunciou aos prazos processuais.')
nota('Mudança de metodologia: a calculadora deixou de somar 0,5% por condição e passou a abrir um campo editável (faixa 0,5%–1,0%) quando há 1 ou 2 condições, travando em 1,5% quando há 3. A mera entrega de documentos exigidos por lei NÃO configura colaboração — conta a utilidade de informações e provas adicionais e inéditas. Admitir o ato sem assumir a responsabilidade jurídica enquadra-se neste inciso (III); o inciso IV exige reconhecimento formal da responsabilidade objetiva.')

heading2('7.4 Art. 23, IV — Admissão Voluntária da Responsabilidade', cor=VERDE_ATEN)
body('Avaliada em dois eixos: conteúdo (parcial/total) × tempestividade (4 momentos processuais).')
tabela(['Condição', '%'],
    [
        ['Sem admissão',                         '0%'],
        ['Parcial — após alegações finais',       '−0,25%'],
        ['Total — após alegações finais',         '−0,5%'],
        ['Parcial — no prazo das alegações finais','−0,5%'],
        ['Total — no prazo das alegações finais', '−1,0%'],
        ['Parcial — no prazo de defesa',          '−1,0%'],
        ['Total — no prazo de defesa',            '−1,5%'],
        ['Parcial — antes da instauração do PAR', '−1,5%'],
        ['Total — antes da instauração do PAR',   '−2,0%'],
    ],
    cor_header=VERDE_ATEN, col_widths=[13, 3])

heading2('7.5 Art. 23, V — Programa de Integridade', cor=VERDE_ATEN)
body('Atenuante de até 5% pela comprovada existência e efetiva aplicação de programa de integridade. O percentual máximo somente é atribuível quando o programa for anterior à prática do ato lesivo (Art. 23, parágrafo único, III).')
cite('Resultado da planilha = [COI × MPI] + APJ   →   se ≥ 1%, aplica-se × 1,25   →   Atenuante = min(resultado, 5%)')
body('Onde: COI = Cultura Organizacional e Instâncias de Governança; MPI = Mecanismos, Políticas e Procedimentos de Integridade; APJ = Atuação da Pessoa Jurídica em relação ao ato lesivo.')
body('Fator multiplicador de 1,25 (Portaria Conjunta CGU nº 6/2022 — Adendo nº 1 ao Manual Prático de Avaliação de PI): incide sobre o resultado da planilha sempre que este for igual ou superior a 1%. O teto de 5% é aplicado DEPOIS do multiplicador.')
tabela(['Resultado da planilha', 'Fator 1,25', 'Resultado final'],
    [
        ['Inferior a 1% (ex.: 0,8%)', 'Não incide', '0,8%'],
        ['Igual ou superior a 1% (ex.: 1,0%)', 'Incide', '1,25%'],
        ['Próximo ao teto (ex.: 4,2%)', 'Incide, mas o teto prevalece', '5,0% (4,2 × 1,25 = 5,25 → teto)'],
    ],
    cor_header=VERDE_ATEN, col_widths=[6, 5, 6])
tabela(['Situação', '%'],
    [['Sem programa ou não comprovado', '0%'], ['Programa apresentado e avaliado (bloco APJ-Anterior)', 'até −5,0%'], ['Programa instituído após o ato lesivo (bloco APJ-Posterior)', 'reduzido']],
    cor_header=VERDE_ATEN, col_widths=[13, 3])
nota('A calculadora exibe o valor pré e pós-multiplicador no painel, no relatório e na minuta, para transparência do cálculo.')

# ══════════════════════════════════════════════════════════════════════════════
#  8. CÁLCULO FINAL E LIMITES LEGAIS
# ══════════════════════════════════════════════════════════════════════════════
heading1('8. CÁLCULO FINAL E LIMITES LEGAIS (Art. 25)')
heading2('8.1 Fórmula da Multa')
cite('Índice de Dosimetria  =  Σ Agravantes (Art. 22)  −  Σ Atenuantes (Art. 23)')
cite('Multa Bruta  =  Faturamento Bruto  ×  Índice de Dosimetria')
heading2('8.2 Verificação dos Limites')
tabela(
    ['Limite', 'Regra', 'Efeito', 'Fundamento'],
    [
        ['Piso — Vantagem Auferida',    'Multa não pode ser inferior à vantagem auferida',        'Se Multa Bruta < VA → Multa = VA', 'Art. 25, I'],
        ['Teto — Maior Vantagem',       'Multa não pode exceder o maior valor entre VA e VP',      'Se Multa Bruta > max(VA,VP) → Multa = max(VA,VP)', 'Art. 25, II'],
        ['Índice ≤ 0',                  'Resultado negativo ou zero na dosimetria',                'Multa = valor mínimo (vantagem auferida)', 'Art. 25, §2º'],
    ],
    col_widths=[4, 5.5, 5, 3]
)

# ══════════════════════════════════════════════════════════════════════════════
#  9. PUBLICAÇÃO EXTRAORDINÁRIA — SEÇÃO APROFUNDADA
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('9. PUBLICAÇÃO EXTRAORDINÁRIA DA DECISÃO CONDENATÓRIA')

heading2('9.1 Natureza Jurídica e Fundamento')
body(
    'A publicação extraordinária da decisão condenatória é uma sanção autônoma prevista no art. 6º, inciso II, '
    'da Lei nº 12.846/2013, aplicada conjuntamente com a multa pecuniária. Não se trata de mera divulgação '
    'de ato processual — constitui penalidade independente, com caráter pedagógico e dissuasório, voltada a '
    'expor publicamente a conduta ilícita da pessoa jurídica e sua responsabilização formal pelo Estado.'
)
cite(
    'Art. 6º, II, Lei nº 12.846/2013. A publicação extraordinária da decisão condenatória ocorrerá na forma '
    'de extrato de sentença, a expensas da pessoa jurídica, em meios de comunicação de grande circulação na '
    'área da prática da infração e de atuação da pessoa jurídica ou, na sua falta, em publicação de circulação '
    'nacional, bem como por meio de afixação de edital no próprio estabelecimento ou no local de exercício da '
    'atividade, de modo visível ao público, e no sítio eletrônico na rede mundial de computadores, respeitando '
    'o prazo determinado nesta dosimetria.'
)
cite(
    'Art. 24, Decreto nº 11.129/2022. A publicação extraordinária da decisão condenatória observará o '
    'prazo fixado a partir da Alíquota de Referência apurada na dosimetria.'
)

heading2('9.2 Modalidades e Meios de Veiculação')
body('A publicação extraordinária deve ser realizada, cumulativamente, nas seguintes formas:')
tabela(
    ['Modalidade', 'Descrição'],
    [
        ['Extrato de sentença em mídia',  'Meios de comunicação de grande circulação na área da prática da infração e de atuação da PJ. Na ausência, publicação de circulação nacional'],
        ['Edital no estabelecimento',     'Afixação em local visível ao público no próprio estabelecimento ou no local de exercício da atividade'],
        ['Sítio eletrônico',              'Publicação no site da PJ na rede mundial de computadores (internet)'],
    ],
    col_widths=[5.5, 11]
)
body('Todos os custos de publicação correm por conta da pessoa jurídica condenada ("a expensas da PJ"), não gerando ônus ao erário.')

heading2('9.3 A Alíquota de Referência')
body(
    'O prazo da publicação extraordinária é dosimetrado a partir de uma grandeza denominada Alíquota de Referência, '
    'que traduz em percentual do faturamento bruto o peso efetivo da sanção aplicada. '
    'Sua apuração segue dois cenários distintos, conforme o mecanismo que determinou o valor final da multa:'
)
heading3('Cenário A — Regra Geral')
body(
    'Aplica-se quando a multa final decorreu diretamente da aplicação da alíquota preliminar resultante da '
    'dosimetria dos arts. 22 e 23, sem interferência dos limites do art. 25 (piso, teto ou vantagem auferida). '
    'Nesse caso, a Alíquota de Referência é a própria alíquota preliminar da dosimetria.'
)
cite('Alíquota de Referência  =  Índice de Dosimetria  =  Σ Agravantes − Σ Atenuantes')

heading3('Cenário B — Regra de Exceção (Limites Legais)')
body(
    'Aplica-se quando o valor final da multa foi alterado pelos limites legais — seja porque a dosimetria resultou '
    'em índice negativo ou zero (art. 25, §2º), seja porque incidiu o piso da vantagem auferida (art. 25, I) ou '
    'o teto da maior vantagem (art. 25, II). Como a alíquota preliminar não corresponde ao valor efetivamente '
    'aplicado, ela é substituída pela alíquota recalculada:'
)
cite('Alíquota de Referência  =  (Valor Final da Multa  ÷  Faturamento Bruto)  ×  100')
nota(
    'O sistema identifica automaticamente o cenário aplicável (A ou B) a partir dos dados inseridos nas etapas '
    'anteriores. A identificação e o cálculo são transparentes e constam do relatório gerado.'
)

heading2('9.4 Tabela de Escalonamento do Prazo')
body(
    'A partir da Alíquota de Referência apurada, o prazo da publicação extraordinária é determinado pela '
    'seguinte tabela de escalonamento, conforme art. 24 do Decreto nº 11.129/2022:'
)
tabela(
    ['Alíquota de Referência', 'Prazo (dias corridos)', 'Prazo por extenso'],
    [
        ['Até 2,5%',                      '30',  'trinta dias'],
        ['Maior que 2,5% até 5,0%',       '45',  'quarenta e cinco dias'],
        ['Maior que 5,0% até 7,5%',       '60',  'sessenta dias'],
        ['Maior que 7,5% até 10,0%',      '75',  'setenta e cinco dias'],
        ['Maior que 10,0% até 12,5%',     '90',  'noventa dias'],
        ['Maior que 12,5% até 15,0%',    '105',  'cento e cinco dias'],
        ['Maior que 15,0% até 17,5%',    '120',  'cento e vinte dias'],
        ['Maior que 17,5%',              '135',  'cento e trinta e cinco dias'],
    ],
    col_widths=[6, 4.5, 6]
)
body('O prazo é contado em dias corridos a partir da publicação da decisão condenatória no órgão oficial correspondente.')

heading2('9.5 Casos Especiais na Apuração da Alíquota de Referência')
heading3('Faturamento igual a zero ou não estimável')
body(
    'Quando a base de cálculo (faturamento bruto) for igual a zero ou não puder ser estimada, o Cenário B '
    'não pode ser aplicado pela impossibilidade aritmética da divisão. Nessa hipótese, adota-se como Alíquota '
    'de Referência o índice da dosimetria preliminar:'
)
cite('Alíquota de Referência  =  max(Índice de Dosimetria, 0)  ×  100')

heading3('Índice de Dosimetria negativo ou zero (art. 25, §2º)')
body(
    'Quando a dosimetria resultar em índice negativo ou zero, a multa é fixada no piso mínimo (vantagem auferida). '
    'Nessa situação, o Cenário B se aplica necessariamente, pois o valor final da multa difere da aplicação '
    'direta da alíquota preliminar. A Alíquota de Referência passa a ser calculada como o quociente entre o '
    'valor da vantagem auferida e o faturamento bruto.'
)

heading3('Vantagem auferida superior à multa bruta (piso)')
body(
    'Quando a vantagem auferida supera a multa bruta calculada pela dosimetria, a multa é elevada ao valor '
    'da vantagem auferida (piso legal, art. 25, I). O Cenário B é aplicado, e a Alíquota de Referência '
    'é recalculada pelo valor efetivamente aplicado.'
)

heading3('Teto da maior vantagem (art. 25, II)')
body(
    'Quando a multa bruta excede o maior valor entre a vantagem auferida e a pretendida, ela é reduzida a '
    'esse teto. O Cenário B é aplicado, e a Alíquota de Referência reflete o valor do teto aplicado.'
)

heading2('9.6 Redação da Decisão — Elementos Obrigatórios')
body('A seção de publicação extraordinária na peça processual deve conter, no mínimo:')
tabela(
    ['Elemento', 'Conteúdo exigido'],
    [
        ['Fundamento legal',         'Art. 6º, II, da Lei nº 12.846/2013 c/c art. 24 do Decreto nº 11.129/2022'],
        ['Cenário adotado',          'A (regra geral) ou B (regra de exceção), com memória de cálculo'],
        ['Alíquota de Referência',   'Percentual apurado, com 2 casas decimais'],
        ['Prazo',                    'Número de dias corridos, por algarismo e por extenso'],
        ['Forma de veiculação',      'Meios de comunicação, edital no estabelecimento e sítio eletrônico'],
        ['Ônus financeiro',          'Expressão de que os custos correm a expensas da pessoa jurídica'],
    ],
    col_widths=[5, 12]
)

heading2('9.7 Geração Automática pela Calculadora')
body(
    'A Etapa 13 da calculadora processa automaticamente todos os dados inseridos nas etapas anteriores e '
    'determina: (i) o cenário aplicável (A ou B); (ii) a Alíquota de Referência; (iii) o prazo correspondente '
    'na tabela de escalonamento; e (iv) o texto dissertativo da memória de cálculo para integrar a minuta '
    'processual.'
)
body('O relatório final impresso registra:')
bullet('Cenário adotado (A ou B) com fundamentação;')
bullet('Memória de cálculo da Alíquota de Referência;')
bullet('Prazo determinado em dias corridos e por extenso;')
bullet('Texto padrão para a cláusula de publicação na decisão condenatória.')
nota(
    'A dosimetria da publicação extraordinária é consequência direta do valor final da multa. Qualquer '
    'alteração nos parâmetros de dosimetria (agravantes, atenuantes, limites) afetará automaticamente o '
    'prazo da publicação. Por isso, a Etapa 13 deve sempre ser revisada após qualquer modificação nas '
    'etapas anteriores.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  10. RELATÓRIO FINAL
# ══════════════════════════════════════════════════════════════════════════════
heading1('10. RELATÓRIO FINAL E IMPRESSÃO')
heading2('10.1 Impressão do Relatório Final (PDF)')
body('O botão "Imprimir Relatório Final (PDF)" aciona o modo de impressão do navegador. O documento gerado tem fundo branco, identidade visual institucional, e reúne todos os parâmetros da dosimetria.')
body('Para gerar o PDF: pressione o botão → selecione "Salvar como PDF" na janela de impressão → ajuste a escala para 80–100% → salve com nome identificador do PAR.')
heading2('10.2 Botões Adicionais')
tabela(['Botão', 'Função'],
    [
        ['Imprimir Avaliação PI',       'Imprime somente a avaliação do programa de integridade (Etapa 12)'],
        ['Gerar Minuta de Relatório',   'Gera texto dissertativo para inserção direta na peça processual'],
        ['Recomeçar',                   'Reinicia o formulário (todos os campos são zerados)'],
    ],
    col_widths=[5.5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
#  11. FLUXO RESUMIDO
# ══════════════════════════════════════════════════════════════════════════════
heading1('11. FLUXO RESUMIDO DO CÁLCULO')
tabela(
    ['Etapa', 'Conteúdo', 'Faixa de %'],
    [
        ['0',  'Faturamento Bruto (base de cálculo)',                     'Base R$'],
        ['1',  'Vantagens Auferida e Pretendida',                         'Piso e Teto'],
        ['2',  'Agrav. I — Concurso de atos lesivos',                     '0% a 4,0%'],
        ['3',  'Agrav. II — Ciência/Tolerância hierárquica',              '0% a 3,0%'],
        ['4',  'Agrav. III — Interrupção/Descumprimento (maior das 3)',   '0% a 4,0%'],
        ['5',  'Agrav. IV — Situação econômica',                          '0% ou 1,0%'],
        ['6',  'Agrav. V — Reincidência',                                 '0% ou 3,0%'],
        ['7',  'Agrav. VI — Contratos com o ente lesado',                 '0% a 5,0%'],
        ['8',  'Aten. I — Não consumação',                                '0% ou −0,5%'],
        ['9',  'Aten. II — Ressarcimento/Devolução',                      '0% a −1,0%'],
        ['10', 'Aten. III — Colaboração',                                 '0% a −1,5%'],
        ['11', 'Aten. IV — Admissão voluntária',                          '0% a −2,0%'],
        ['12', 'Aten. V — Programa de integridade',                       '0% a −5,0%'],
        ['13', 'Publicação Extraordinária + Resultado Final',             'Prazo 30–135 dias'],
    ],
    col_widths=[1.5, 10, 5]
)

# ══════════════════════════════════════════════════════════════════════════════
#  12. FAQ
# ══════════════════════════════════════════════════════════════════════════════
heading1('12. PERGUNTAS FREQUENTES E PONTOS DE ATENÇÃO')

heading2('O sistema é validado por órgão oficial?')
body('A calculadora foi desenvolvida pela Corregedoria da RFB como ferramenta de apoio. Os parâmetros refletem o Decreto nº 11.129/2022, a IN CGU nº 1/2015 e o Guia CGU. O resultado é uma proposta fundamentada, sujeita à deliberação da autoridade competente.')

heading2('Como tratar a atualização pelo IPCA?')
body('A calculadora não realiza atualização automática pelo IPCA por funcionar offline. Quando necessário (Art. 21), corrija o valor previamente na Calculadora do Cidadão (BCB/IPCA) e informe o resultado corrigido no campo correspondente.')

heading2('O que acontece se o índice de dosimetria for negativo?')
body('Nos termos do Art. 25, §2º do Decreto nº 11.129/2022, a multa é fixada no valor mínimo (vantagem auferida). Se a vantagem auferida for zero ou não estimável, aplica-se o piso legal de R$ 6 mil nas hipóteses cabíveis.')

heading2('Qual a diferença entre vantagem auferida e pretendida na prática?')
body('Auferida: valor efetivamente recebido, economizado ou patrimonializado pela PJ. Pretendida: valor planejado que seria obtido se o ato ilícito fosse plenamente executado — inferido de proposta, contrato ou estimativa.')

heading2('O programa de integridade posterior ao ato lesivo pode atenuar?')
body('Sim, mas com percentual reduzido. O bloco APJ-Posterior avalia medidas corretivas adotadas após a descoberta. O percentual máximo de 5% é exclusivo de programas preexistentes ao ato lesivo.')

heading2('Quantos dias tem a publicação extraordinária nos casos mais graves?')
body('O prazo máximo é de 135 dias corridos, aplicável quando a Alíquota de Referência superar 17,5%. Para alíquotas até 2,5%, o prazo mínimo é de 30 dias.')

heading2('O Art. 22, IV se aplica quando a Liquidez Geral é 0,98 mas os demais indicadores são favoráveis?')
body('Não. O agravante de situação econômica exige os três indicadores cumulativamente: Solvência Geral > 1, Liquidez Geral > 1 e Lucro Líquido positivo. Qualquer indicador abaixo do limiar — ainda que por margem ínfima — zera o agravante.')

heading2('Responder a outro PAR em andamento configura reincidência?')
body('Não. A reincidência (Art. 22, V) exige decisão definitiva anterior. Estar respondendo a outro PAR sem julgamento concluído não configura reincidência, pois o marco legal é a publicação do julgamento anterior, não a instauração de novo processo.')

heading2('Como é calculada a publicação quando a multa é ajustada pelo piso legal?')
body('Quando a multa final é elevada ao valor da vantagem auferida (piso), aplica-se o Cenário B: Alíquota de Referência = (Multa Final ÷ Faturamento Bruto) × 100. Isso garante que o prazo de publicação reflita a gravidade proporcional real da sanção.')

heading2('Como a colaboração é valorada com 1, 2 ou 3 condições presentes? (atualização CGU 2026)')
body('Conforme a Tabela 6 do Manual CGU, uma ou duas condições ensejam percentual discricionário de 0,5% a 1,0% (fixado pela autoridade conforme a utilidade da colaboração); as três condições, simultaneamente, ensejam 1,5% fixo. A calculadora abre um campo editável (faixa 0,5–1,0%) para 1 ou 2 condições e trava em 1,5% para 3 — não soma mais 0,5% por condição.')

heading2('O que é o fator de 1,25 no Programa de Integridade? (atualização CGU 2026)')
body('Conforme a Portaria Conjunta CGU nº 6/2022, ao resultado da planilha [COI × MPI] + APJ aplica-se um fator multiplicador de 1,25 sempre que esse resultado for igual ou superior a 1%, respeitado o teto de 5%. O fator não incide sobre resultados abaixo de 1%, e o teto de 5% é aplicado depois do multiplicador (ex.: 4,2% × 1,25 = 5,25% → limitado a 5,0%).')

# ══════════════════════════════════════════════════════════════════════════════
#  13. CASOS ESPECIAIS
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('13. CASOS ESPECIAIS E SITUAÇÕES ATÍPICAS')

heading2('13.1 Empresa sem Defesa Escrita ou sem Balanço Apresentado')
body('Quando a PJ não apresenta defesa escrita nem documentação contábil, o operador deve buscar dados em sistemas internos antes de qualquer outra medida:')
tabela(['Sistema / Fonte', 'Informação disponível'],
    [
        ['ECF (Escrituração Contábil Fiscal)', 'Principal fonte para Lucro Real/Presumido — DRE e Balanço'],
        ['DIRF / DCTF',                        'Cruzamento de rendimentos e declarações'],
        ['SPED',                               'Escrituração contábil e fiscal digital'],
        ['DAS / PGDAS',                        'Faturamento para optantes do Simples Nacional'],
        ['NF-e / NFS-e',                       'Notas fiscais emitidas — úteis para estimativa'],
    ],
    col_widths=[6, 11])
nota('A ausência de defesa não impede a dosimetria nem desobriga o operador de buscar exaustivamente os dados antes de recorrer à estimativa.')

heading2('13.2 Piso Absoluto de R$ 6.000,00 (Empresa Fantasma ou Noteira)')
body('Quando a PJ é uma empresa-fantasma sem faturamento rastreável e a alíquota resultante é zero ou ínfima, aplica-se o piso absoluto de R$ 6.000,00 (Art. 20, §4º da Lei nº 12.846/2013). A multa final não pode ser inferior a esse valor.')
tabela(['Situação', 'Multa Final'],
    [
        ['Alíquota zero por excesso de atenuantes', 'R$ 6.000,00 (piso absoluto)'],
        ['Faturamento ínfimo + alíquota baixa → resultado < R$ 6k', 'R$ 6.000,00 (piso absoluto)'],
    ],
    col_widths=[10, 7])

heading2('13.3 Estimativa Fundamentada (Art. 20, §1º, III)')
body('Quando não há faturamento histórico em nenhum sistema, a autoridade realiza estimativa técnica com base em elementos objetivos. O procedimento é:')
bullet('1. Levantar o capital social ou outros indicadores objetivos (GFIP, compras, valor de repasses ilícitos);')
bullet('2. Identificar o CNAE e a carga tributária estimada do regime vigente;')
bullet('3. Base histórica líquida = indicador bruto − tributos estimados;')
bullet('4. Atualizar pelo IPCA até 31/12 do exercício anterior ao PAR.')
nota('Estimativa técnica não é valor arbitrário. Deve ser formalmente fundamentada com os elementos usados. A autoridade documenta cada premissa no processo.')

heading2('13.4 Retroação Anual — Art. 21 (Sem Faturamento no Exercício Anterior)')
body('Quando a PJ não teve faturamento no exercício imediatamente anterior ao PAR, recua-se ano a ano até encontrar um período com faturamento válido, e atualiza-se pelo IPCA:')
tabela(['Exercício', 'Faturamento Bruto', 'Situação'],
    [
        ['Exercício PAR − 1',  'R$ 0,00',         'Sem faturamento — retroage'],
        ['Exercício PAR − 2',  'R$ 0,00',         'Sem faturamento — retroage'],
        ['Exercício PAR − N',  'Valor encontrado', 'Base histórica → aplicar IPCA'],
    ],
    col_widths=[5, 5, 7])
body('O valor corrigido pelo IPCA passa a ser a base de cálculo oficial para toda a dosimetria.')

heading2('13.5 Sucessão Empresarial (Art. 4º, Lei nº 12.846/2013)')
tabela(['Evento Societário', 'Responsabilidade da Sucessora', 'Base de Cálculo'],
    [
        ['Alteração contratual ou transformação', 'Mesma PJ — responsabilidade inalterada', 'Faturamento da própria PJ'],
        ['Incorporação ou fusão', 'Sucessora responde, limitada ao patrimônio transferido', 'Faturamento da empresa SUCESSORA'],
    ],
    col_widths=[5.5, 6, 5.5])
nota('Se o evento sucessório ocorrer durante o andamento do PAR, o polo passivo é redirecionado na Nota de Indiciação para a empresa sucessora.')

heading2('13.6 Grupo Econômico (Art. 20, §2º, Decreto nº 11.129/2022)')
body('Quando há prova de que diferentes empresas agiram em conluio ou fraude estruturada como se fossem uma única entidade, configura-se o grupo econômico para fins de dosimetria. Consequências:')
bullet('Consolidação dos faturamentos brutos de todas as PJs participantes para compor a base de cálculo;')
bullet('Avaliação conjunta de agravantes e atenuantes;')
bullet('Consideração da vantagem auferida por todo o esquema ilícito;')
bullet('Solidariedade: todas as PJs integrantes respondem integralmente pelo valor total da multa aplicada.')

# ══════════════════════════════════════════════════════════════════════════════
#  14. ESTUDOS DE CASO
# ══════════════════════════════════════════════════════════════════════════════
page_break()
heading1('14. ESTUDOS DE CASO COMENTADOS')
body('Os casos a seguir ilustram a aplicação concreta da dosimetria, com dados numéricos reais e as decisões metodológicas tomadas em cada etapa.')

heading2('14.1 Caso I — Empresa Alimentos Sabor SA: Fraude em Fiscalização Tributária')
heading3('Cenário')
body('PAR instaurado em 2025. CNPJ 12.345.678/0001-90. Ato lesivo: fraude em fiscalização tributária. Provas: Operação Policial com quebra de sigilos telemático, telefônico e bancário. Benefício econômico estimado: R$ 90.000.000,00.')

heading3('Passo 1 — Base de Cálculo (ECF/DRE da Matriz, exercício 2024)')
tabela(['Item', 'Valor'],
    [
        ['Faturamento Bruto', 'R$ 1.248.765.432,00'],
        ['(−) Tributos Incidentes', '(R$ 188.653.322,00)'],
        ['Base de Cálculo Final', 'R$ 1.060.112.110,00'],
    ],
    col_widths=[9, 8])

heading3('Passo 2 — Agravantes e Atenuantes')
tabela(['Circunstância', 'Fundamentação', '%'],
    [
        ['Art. 22, I — Concurso de atos', 'Múltiplos pagamentos dissimulados (R$ 450k via empresa B, R$ 250k via empresa C, R$ 180k em móveis para o auditor)', '3,5%'],
        ['Art. 22, II — Ciência do corpo diretivo', 'WhatsApp do gerente citando o Administrador; e-mails copiados ao contador e auditor → ciência do Administrador (teto)', '3,0%'],
        ['Art. 22, III — Interrupção', 'Não aplicável ao caso', '0%'],
        ['Art. 22, IV — Situação econômica', 'SG = 1,67 ✓ / LG = 0,98 ✗ / Lucro ✓ — LG < 1, não se aplica', '0%'],
        ['Art. 22, V — Reincidência', 'PAR anterior em andamento sem julgamento definitivo — NÃO configura reincidência', '0%'],
        ['Art. 22, VI — Contratos', 'Não pontuado no caso', '0%'],
        ['TOTAL AGRAVANTES', '', '6,5%'],
        ['Art. 23, II — Ressarcimento', 'Comprovação parcial de devolução espontânea', '−0,5%'],
        ['TOTAL ATENUANTES', '', '−0,5%'],
        ['ALÍQUOTA FINAL', '6,5% − 0,5%', '6,0%'],
    ],
    col_widths=[5, 9.5, 2.5])

heading3('Passo 3 — Multa Preliminar')
cite('R$ 1.060.112.110,00 × 6,0%  =  R$ 63.606.726,60')

heading3('Passo 4 — Limites Legais')
tabela(['Limite', 'Fórmula', 'Valor'],
    [
        ['Mínimo (maior entre os dois)', 'Vantagem Auferida (após refiscalização)', 'R$ 60.000.000,00'],
        ['', '0,1% do Faturamento', 'R$ 1.060.112,11'],
        ['Máximo (menor entre os dois)', '20% do Faturamento', 'R$ 212.022.422,00'],
        ['', '3 × Vantagem Pretendida (R$ 90M)', 'R$ 270.000.000,00'],
    ],
    col_widths=[5.5, 7, 5])

heading3('Passo 5 — Multa Final (Caso I)')
body('A Multa Preliminar de R$ 63.606.726,60 está entre o piso (R$ 60M) e o teto (R$ 212M) — não sofre ajuste.')
cite('MULTA FINAL (Caso I)  =  R$ 63.606.726,60')

heading3('Passo 6 — Publicação Extraordinária')
body('Multa não foi limitada (Cenário A). Alíquota de referência = 6,0% → faixa "maior que 5,0% até 7,5%" → PRAZO: 60 DIAS.')

heading2('14.2 Caso II — Mesmo Cenário, Vantagem Auferida Integral (Piso Forçado)')
heading3('Variante')
body('A vantagem auferida permanece em R$ 90.000.000,00 (sem redução pela refiscalização).')

tabela(['Limite', 'Valor'],
    [
        ['Piso (vantagem auferida integral)', 'R$ 90.000.000,00'],
        ['Teto (20% do Faturamento)', 'R$ 212.022.422,00'],
    ],
    col_widths=[9, 8])

body('Multa Preliminar de R$ 63.606.726,60 ficou abaixo do piso legal (R$ 90M). A multa é elevada obrigatoriamente ao piso:')
cite('MULTA FINAL (Caso II)  =  R$ 90.000.000,00')

heading3('Publicação com Proporção Recalculada (Cenário B)')
cite('Alíquota de Proporção  =  R$ 90.000.000,00 ÷ R$ 1.060.112.110,00  ≈  8,5%  →  PRAZO: 75 DIAS')
nota('O prazo de publicação aumentou de 60 para 75 dias porque a gravidade proporcional da sanção cresceu quando a multa foi elevada ao piso.')

heading2('14.3 Caso III — Empresa Cítrica: Sem Defesa, Empresa Fantasma')
heading3('Cenário')
body('Empresa-noteira que não apresentou defesa nem balanço. Faturamento: 2024 e 2025 = R$ 0. Faturamento histórico: 2018 (jan–mar) = R$ 53.500,00; 2017 = R$ 560.000,00.')

heading3('Retroação e Atualização (Art. 21)')
cite('Base de Cálculo apurada e atualizada pelo IPCA até 31/12/2024  =  R$ 74.588,17')

heading3('Dosimetria')
tabela(['Circunstância', '%'],
    [
        ['Agravantes', '0%'],
        ['Art. 23, II — Inexistência de vantagem/dano comprovado', '−1,0%'],
        ['Alíquota Final (atenuantes > agravantes)', '0%'],
    ],
    col_widths=[13, 4])

cite('Multa Preliminar  =  R$ 74.588,17 × 0%  =  R$ 0,00')
body('Piso absoluto (Art. 20, §4º, Lei nº 12.846/2013):')
cite('MULTA FINAL (Caso III)  =  R$ 6.000,00')
cite('Alíquota de Proporção  =  R$ 6.000,00 ÷ R$ 74.588,17  ≈  8,0%  →  PRAZO: 75 DIAS')

heading2('14.4 Caso IV — Empresa Gasparzinho: Estimativa Fundamentada')
heading3('Cenário')
body('Empresa sem qualquer dado de faturamento histórico. Capital social (2017): R$ 50.000,00. CNAE: 6204-0/00 (Consultoria em TI). Sede: Não-Me-Toque/RS. Regime: Lucro Presumido.')

tabela(['Item', 'Valor'],
    [
        ['Capital Social 2017 (proxy de faturamento)', 'R$ 50.000,00'],
        ['(−) Tributos estimados: PIS 0,65% + COFINS 3% + ISS 3% = 6,65%', '(R$ 3.325,00)'],
        ['Base de Cálculo Histórica (2017)', 'R$ 46.675,00'],
        ['Atualização IPCA até dez/2024', '≈ R$ 67.417,43'],
    ],
    col_widths=[11, 6])

cite('BASE DE CÁLCULO OFICIAL  =  R$ 67.417,43')
body('A partir desse valor, a dosimetria (agravantes, atenuantes, limites legais) segue o fluxo padrão.')
nota('Estimativa técnica não é arbitrária. Cada premissa (CNAE, carga tributária, capital social) deve ser formalmente documentada no processo.')

# ── Rodapé ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p_rodape = doc.add_paragraph()
p_rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rod = p_rodape.add_run(
    'Manual elaborado com base na Calculadora de Multa PAR — Corregedoria da Receita Federal do Brasil\n'
    'Fundamentos: Lei nº 12.846/2013 | Decreto nº 11.129/2022 | IN CGU nº 1/2015 | Manual CGU (3ª ed., jun/2026) | Portaria Conjunta CGU nº 6/2022'
)
r_rod.font.name = 'Arial'; r_rod.font.size = Pt(9)
r_rod.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r_rod.font.italic = True

# ── Salvar ────────────────────────────────────────────────────────────────────
output_path = '/home/user/desktop-tutorial/Manual_Calculadora_PAR.docx'
doc.save(output_path)
print(f'Salvo em: {output_path}')
